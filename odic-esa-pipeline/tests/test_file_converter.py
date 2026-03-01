"""
Tests for the FileConverter skill.

Tests conversion of various file types to PDF:
- Word documents (.docx)
- Images (.jpg, .png, .tiff)
- PDF passthrough
- Multiple image combination
"""

import asyncio
import os
import tempfile
from pathlib import Path
from io import BytesIO

import pytest
from PIL import Image
from docx import Document
from PyPDF2 import PdfReader

# Add parent to path for imports
import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from skills.file_converter import FileConverter


@pytest.fixture
def config():
    """Basic configuration for tests."""
    return {
        'converter': {
            'output_dir': tempfile.mkdtemp(),
            'image_dpi': 150,
            'image_quality': 85,
            'max_image_dimension': 2000,
            'page_size': 'letter',
            'margin_inches': 0.5,
        },
        'pipeline': {
            'staging_dir': tempfile.mkdtemp(),
        }
    }


@pytest.fixture
def converter(config):
    """Create a FileConverter instance."""
    return FileConverter(config)


@pytest.fixture
def temp_dir():
    """Create a temporary directory for test files."""
    import shutil
    d = tempfile.mkdtemp()
    yield d
    shutil.rmtree(d, ignore_errors=True)


# ===== Helper Functions =====

def create_test_word_doc(path: Path, content: str = "Test document content."):
    """Create a simple Word document for testing."""
    doc = Document()
    doc.add_heading('Test Document', 0)
    for line in content.split('\n'):
        doc.add_paragraph(line)
    doc.save(str(path))
    return path


def create_test_image(path: Path, size=(800, 600), color='red', format='JPEG'):
    """Create a simple test image."""
    img = Image.new('RGB', size, color)
    img.save(str(path), format=format)
    return path


def create_test_pdf(path: Path, pages=1):
    """Create a simple test PDF."""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter

    c = canvas.Canvas(str(path), pagesize=letter)
    for i in range(pages):
        c.drawString(100, 700, f"Test page {i + 1}")
        if i < pages - 1:
            c.showPage()
    c.save()
    return path


# ===== Basic Tests =====

class TestFileConverterInit:
    """Tests for FileConverter initialization."""

    def test_init_with_config(self, config):
        """Test converter initializes with config."""
        converter = FileConverter(config)
        assert converter is not None
        assert converter.output_dir.exists()

    def test_default_settings(self, config):
        """Test default settings are applied."""
        converter = FileConverter(config)
        assert converter.image_dpi == 150
        assert converter.image_quality == 85

    def test_supported_extensions(self, converter):
        """Test supported extensions are returned."""
        extensions = converter.get_supported_extensions()
        assert '.pdf' in extensions
        assert '.docx' in extensions
        assert '.jpg' in extensions
        assert '.png' in extensions
        assert '.tiff' in extensions


class TestInputValidation:
    """Tests for input validation."""

    def test_validate_existing_pdf(self, converter, temp_dir):
        """Test validation accepts existing PDF."""
        pdf_path = create_test_pdf(Path(temp_dir) / 'test.pdf')
        assert converter.validate_input(str(pdf_path)) is True

    def test_validate_existing_docx(self, converter, temp_dir):
        """Test validation accepts existing DOCX."""
        docx_path = create_test_word_doc(Path(temp_dir) / 'test.docx')
        assert converter.validate_input(str(docx_path)) is True

    def test_validate_existing_image(self, converter, temp_dir):
        """Test validation accepts existing images."""
        jpg_path = create_test_image(Path(temp_dir) / 'test.jpg')
        assert converter.validate_input(str(jpg_path)) is True

    def test_validate_nonexistent_file(self, converter):
        """Test validation rejects nonexistent file."""
        assert converter.validate_input('/nonexistent/file.pdf') is False

    def test_validate_unsupported_extension(self, converter, temp_dir):
        """Test validation rejects unsupported extension."""
        txt_path = Path(temp_dir) / 'test.txt'
        txt_path.write_text('test')
        assert converter.validate_input(str(txt_path)) is False

    def test_validate_list_of_files(self, converter, temp_dir):
        """Test validation accepts list of valid files."""
        img1 = create_test_image(Path(temp_dir) / 'img1.jpg')
        img2 = create_test_image(Path(temp_dir) / 'img2.jpg')
        assert converter.validate_input([str(img1), str(img2)]) is True


# ===== PDF Passthrough Tests =====

class TestPDFPassthrough:
    """Tests for PDF passthrough (no conversion needed)."""

    @pytest.mark.asyncio
    async def test_pdf_passthrough(self, converter, temp_dir):
        """Test that PDFs pass through without conversion."""
        pdf_path = create_test_pdf(Path(temp_dir) / 'test.pdf')

        result = await converter.process(str(pdf_path))

        assert result.success is True
        assert result.data['conversion_type'] == 'passthrough'
        assert result.data['output_path'] == str(pdf_path)

    @pytest.mark.asyncio
    async def test_pdf_preserves_original(self, converter, temp_dir):
        """Test that original PDF is not modified."""
        pdf_path = create_test_pdf(Path(temp_dir) / 'test.pdf')
        original_size = pdf_path.stat().st_size

        await converter.process(str(pdf_path))

        assert pdf_path.stat().st_size == original_size


# ===== Word Document Conversion Tests =====

class TestWordConversion:
    """Tests for Word document to PDF conversion."""

    @pytest.mark.asyncio
    async def test_word_to_pdf_basic(self, converter, temp_dir):
        """Test basic Word to PDF conversion."""
        docx_path = create_test_word_doc(
            Path(temp_dir) / 'test.docx',
            "This is a test document.\nWith multiple lines."
        )

        result = await converter.process(str(docx_path))

        assert result.success is True
        assert result.data['conversion_type'] == 'word_to_pdf'
        assert Path(result.data['output_path']).exists()
        assert Path(result.data['output_path']).suffix == '.pdf'

    @pytest.mark.asyncio
    async def test_word_creates_valid_pdf(self, converter, temp_dir):
        """Test that converted Word doc creates valid, readable PDF."""
        docx_path = create_test_word_doc(Path(temp_dir) / 'test.docx', "Test content")

        result = await converter.process(str(docx_path))

        # Verify PDF is readable
        reader = PdfReader(result.data['output_path'])
        assert len(reader.pages) >= 1

    @pytest.mark.asyncio
    async def test_word_with_long_content(self, converter, temp_dir):
        """Test Word doc with content spanning multiple pages."""
        long_content = '\n'.join(['This is paragraph number ' + str(i) for i in range(100)])
        docx_path = create_test_word_doc(Path(temp_dir) / 'long.docx', long_content)

        result = await converter.process(str(docx_path))

        assert result.success is True
        assert result.data.get('pages_created', 1) >= 1

    @pytest.mark.asyncio
    async def test_word_metadata_extraction(self, converter, temp_dir):
        """Test that conversion captures metadata."""
        docx_path = create_test_word_doc(Path(temp_dir) / 'test.docx')

        result = await converter.process(str(docx_path))

        assert 'paragraphs_converted' in result.data
        assert 'pages_created' in result.data


# ===== Image Conversion Tests =====

class TestImageConversion:
    """Tests for image to PDF conversion."""

    @pytest.mark.asyncio
    async def test_jpg_to_pdf(self, converter, temp_dir):
        """Test JPEG to PDF conversion."""
        jpg_path = create_test_image(Path(temp_dir) / 'test.jpg', format='JPEG')

        result = await converter.process(str(jpg_path))

        assert result.success is True
        assert result.data['conversion_type'] == 'image_to_pdf'
        assert Path(result.data['output_path']).exists()

    @pytest.mark.asyncio
    async def test_png_to_pdf(self, converter, temp_dir):
        """Test PNG to PDF conversion."""
        png_path = create_test_image(Path(temp_dir) / 'test.png', format='PNG')

        result = await converter.process(str(png_path))

        assert result.success is True
        assert result.data['conversion_type'] == 'image_to_pdf'

    @pytest.mark.asyncio
    async def test_tiff_to_pdf(self, converter, temp_dir):
        """Test TIFF to PDF conversion."""
        tiff_path = create_test_image(Path(temp_dir) / 'test.tiff', format='TIFF')

        result = await converter.process(str(tiff_path))

        assert result.success is True

    @pytest.mark.asyncio
    async def test_png_with_transparency(self, converter, temp_dir):
        """Test PNG with transparency converts correctly."""
        img_path = Path(temp_dir) / 'transparent.png'
        img = Image.new('RGBA', (400, 300), (255, 0, 0, 128))  # Semi-transparent red
        img.save(str(img_path), format='PNG')

        result = await converter.process(str(img_path))

        assert result.success is True
        # Should convert to RGB (no transparency in PDF)
        reader = PdfReader(result.data['output_path'])
        assert len(reader.pages) == 1

    @pytest.mark.asyncio
    async def test_large_image_scaled(self, converter, temp_dir):
        """Test that large images are scaled down."""
        # Create image larger than max_image_dimension (2000 in config)
        large_path = create_test_image(Path(temp_dir) / 'large.jpg', size=(5000, 4000))

        result = await converter.process(str(large_path))

        assert result.success is True
        assert result.data.get('scale_factor', 1.0) < 1.0

    @pytest.mark.asyncio
    async def test_image_metadata(self, converter, temp_dir):
        """Test image conversion captures metadata."""
        jpg_path = create_test_image(Path(temp_dir) / 'test.jpg', size=(800, 600))

        result = await converter.process(str(jpg_path))

        assert 'original_size' in result.data
        assert result.data['original_size'] == (800, 600)


# ===== Multi-Image Combination Tests =====

class TestMultiImageConversion:
    """Tests for combining multiple images into a single PDF."""

    @pytest.mark.asyncio
    async def test_multiple_images_to_pdf(self, converter, temp_dir):
        """Test combining multiple images into multi-page PDF."""
        images = [
            create_test_image(Path(temp_dir) / f'img{i}.jpg', color=c)
            for i, c in enumerate(['red', 'green', 'blue'])
        ]

        result = await converter.process([str(p) for p in images])

        assert result.success is True
        assert result.data['conversion_type'] == 'images_to_multi_page_pdf'
        assert result.data['images_converted'] == 3
        assert result.data['total_pages'] == 3

    @pytest.mark.asyncio
    async def test_multi_image_creates_valid_pdf(self, converter, temp_dir):
        """Test that multi-image PDF is valid and has correct page count."""
        images = [
            create_test_image(Path(temp_dir) / f'img{i}.jpg')
            for i in range(4)
        ]

        result = await converter.process([str(p) for p in images])

        reader = PdfReader(result.data['output_path'])
        assert len(reader.pages) == 4

    @pytest.mark.asyncio
    async def test_mixed_image_formats(self, converter, temp_dir):
        """Test combining images of different formats."""
        jpg = create_test_image(Path(temp_dir) / 'test.jpg', format='JPEG')
        png = create_test_image(Path(temp_dir) / 'test.png', format='PNG')

        result = await converter.process([str(jpg), str(png)])

        assert result.success is True
        assert result.data['images_converted'] == 2


# ===== Batch Processing Tests =====

class TestBatchProcessing:
    """Tests for batch processing of mixed file types."""

    @pytest.mark.asyncio
    async def test_batch_mixed_types(self, converter, temp_dir):
        """Test batch processing of mixed file types."""
        files = [
            create_test_word_doc(Path(temp_dir) / 'doc.docx'),
            create_test_image(Path(temp_dir) / 'img.jpg'),
            create_test_pdf(Path(temp_dir) / 'test.pdf'),
        ]

        result = await converter.process([str(f) for f in files])

        assert result.success is True
        assert result.data['conversion_type'] == 'batch'
        assert result.data['total_files'] == 3
        assert result.data['successful'] == 3
        assert result.data['failed'] == 0

    @pytest.mark.asyncio
    async def test_batch_with_failures(self, converter, temp_dir):
        """Test batch processing handles failures gracefully with mixed file types."""
        # Use mixed file types to trigger individual processing (not multi-image)
        good_image = create_test_image(Path(temp_dir) / 'good.jpg')
        good_doc = create_test_word_doc(Path(temp_dir) / 'good.docx')
        bad_file = Path(temp_dir) / 'bad.txt'
        bad_file.write_text('not an image')
        # Rename to trick validation but fail conversion
        bad_jpg = bad_file.rename(Path(temp_dir) / 'bad.jpg')

        result = await converter.process([str(good_image), str(good_doc), str(bad_jpg)])

        assert result.success is True  # Batch still succeeds
        assert result.data['conversion_type'] == 'batch'
        assert result.data['successful'] >= 2  # At least the good files
        # At least one failure expected
        assert any(not r['success'] for r in result.data['results'])


# ===== Synchronous Interface Tests =====

class TestSyncInterface:
    """Tests for the synchronous conversion interface."""

    def test_sync_word_conversion(self, converter, temp_dir):
        """Test synchronous Word conversion."""
        docx_path = create_test_word_doc(Path(temp_dir) / 'test.docx')

        success, output_path, metadata = converter.convert_sync(docx_path)

        assert success is True
        assert Path(output_path).exists()
        assert 'conversion_type' in metadata

    def test_sync_image_conversion(self, converter, temp_dir):
        """Test synchronous image conversion."""
        jpg_path = create_test_image(Path(temp_dir) / 'test.jpg')

        success, output_path, metadata = converter.convert_sync(jpg_path)

        assert success is True
        assert Path(output_path).exists()


# ===== Edge Cases =====

class TestEdgeCases:
    """Tests for edge cases and error handling."""

    @pytest.mark.asyncio
    async def test_empty_word_doc(self, converter, temp_dir):
        """Test handling of empty Word document."""
        docx_path = Path(temp_dir) / 'empty.docx'
        doc = Document()  # Empty document
        doc.save(str(docx_path))

        result = await converter.process(str(docx_path))

        # Should still succeed, just with minimal content
        assert result.success is True

    @pytest.mark.asyncio
    async def test_tiny_image(self, converter, temp_dir):
        """Test handling of very small image."""
        tiny_path = create_test_image(Path(temp_dir) / 'tiny.jpg', size=(1, 1))

        result = await converter.process(str(tiny_path))

        assert result.success is True

    @pytest.mark.asyncio
    async def test_grayscale_image(self, converter, temp_dir):
        """Test handling of grayscale image."""
        img_path = Path(temp_dir) / 'gray.jpg'
        img = Image.new('L', (400, 300), 128)  # Grayscale
        img.save(str(img_path), format='JPEG')

        result = await converter.process(str(img_path))

        assert result.success is True

    @pytest.mark.asyncio
    async def test_output_path_uniqueness(self, converter, temp_dir):
        """Test that output paths are unique for same input."""
        jpg_path = create_test_image(Path(temp_dir) / 'test.jpg')

        result1 = await converter.process(str(jpg_path))
        result2 = await converter.process(str(jpg_path))

        # Microseconds in timestamp ensure uniqueness
        assert result1.data['output_path'] != result2.data['output_path']


# ===== Integration Tests =====

class TestFileTypeDetection:
    """Tests for file type detection."""

    def test_detect_word(self, converter, temp_dir):
        """Test Word document detection."""
        docx_path = create_test_word_doc(Path(temp_dir) / 'test.docx')
        file_type = converter._get_file_type(docx_path)
        assert file_type == 'word'

    def test_detect_images(self, converter, temp_dir):
        """Test image detection for various formats."""
        for ext, fmt in [('.jpg', 'JPEG'), ('.png', 'PNG'), ('.tiff', 'TIFF')]:
            img_path = create_test_image(Path(temp_dir) / f'test{ext}', format=fmt)
            file_type = converter._get_file_type(img_path)
            assert file_type == 'image', f"Failed for {ext}"

    def test_detect_pdf(self, converter, temp_dir):
        """Test PDF detection."""
        pdf_path = create_test_pdf(Path(temp_dir) / 'test.pdf')
        file_type = converter._get_file_type(pdf_path)
        assert file_type == 'pdf'

    def test_detect_unsupported(self, converter, temp_dir):
        """Test unsupported file type returns None."""
        txt_path = Path(temp_dir) / 'test.txt'
        txt_path.write_text('test')
        file_type = converter._get_file_type(txt_path)
        assert file_type is None
