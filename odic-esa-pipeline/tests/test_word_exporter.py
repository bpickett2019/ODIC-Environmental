"""
Tests for the Word Exporter skill.

Tests PDF to Word document conversion functionality including:
- Basic PDF conversion
- Fallback text extraction mode
- Batch export
- Error handling
"""

import asyncio
import os
import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

# Test fixtures and helper functions
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


def create_test_pdf(path: Path, num_pages: int = 3, with_text: bool = True) -> Path:
    """Create a test PDF file with the given number of pages."""
    c = canvas.Canvas(str(path), pagesize=letter)

    for i in range(num_pages):
        if with_text:
            c.setFont("Helvetica", 12)
            c.drawString(100, 700, f"Test Page {i + 1}")
            c.drawString(100, 680, f"This is test content for page {i + 1}")
            c.drawString(100, 660, "ODIC Environmental - Phase I ESA")
            c.drawString(100, 640, "1.0 Introduction")
            c.drawString(100, 620, "This section covers the introduction.")
        c.showPage()

    c.save()
    return path


def create_test_word_doc(path: Path, paragraphs: int = 5) -> Path:
    """Create a test Word document."""
    doc = Document()
    doc.add_heading('Test Document', level=1)

    for i in range(paragraphs):
        doc.add_paragraph(f'This is paragraph {i + 1} of the test document.')

    doc.save(str(path))
    return path


# ===== Fixtures =====

@pytest.fixture
def config():
    """Standard test configuration."""
    return {
        'word_exporter': {
            'output_dir': tempfile.mkdtemp(),
            'preserve_images': True,
            'preserve_tables': True
        },
        'pipeline': {
            'output_dir': tempfile.mkdtemp()
        },
        'debug': True
    }


@pytest.fixture
def temp_dir():
    """Create a temporary directory for test files."""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield tmpdir


@pytest.fixture
def exporter(config):
    """Create a WordExporter instance."""
    from skills.word_exporter import WordExporter
    return WordExporter(config)


# ===== Initialization Tests =====

class TestWordExporterInit:
    """Tests for WordExporter initialization."""

    def test_init_with_config(self, config):
        """Test initialization with configuration."""
        from skills.word_exporter import WordExporter
        exporter = WordExporter(config)

        assert exporter.config == config
        assert exporter.output_dir.exists()

    def test_default_output_dir(self):
        """Test that default output directory is created."""
        from skills.word_exporter import WordExporter
        config = {'pipeline': {'output_dir': tempfile.mkdtemp()}}
        exporter = WordExporter(config)

        assert exporter.output_dir.exists()

    def test_pdf2docx_check(self, config):
        """Test pdf2docx availability check."""
        from skills.word_exporter import WordExporter
        exporter = WordExporter(config)

        # Should be True if pdf2docx is installed, False otherwise
        assert isinstance(exporter._pdf2docx_available, bool)


# ===== Input Validation Tests =====

class TestInputValidation:
    """Tests for input validation."""

    def test_validate_existing_pdf(self, exporter, temp_dir):
        """Test validation of existing PDF file."""
        pdf_path = create_test_pdf(Path(temp_dir) / 'test.pdf')
        assert exporter.validate_input(str(pdf_path)) is True

    def test_validate_dict_input(self, exporter, temp_dir):
        """Test validation of dict input with pdf_path."""
        pdf_path = create_test_pdf(Path(temp_dir) / 'test.pdf')
        input_data = {'pdf_path': str(pdf_path)}
        assert exporter.validate_input(input_data) is True

    def test_validate_dict_with_path_key(self, exporter, temp_dir):
        """Test validation of dict input with 'path' key."""
        pdf_path = create_test_pdf(Path(temp_dir) / 'test.pdf')
        input_data = {'path': str(pdf_path)}
        assert exporter.validate_input(input_data) is True

    def test_validate_nonexistent_file(self, exporter):
        """Test validation of nonexistent file."""
        assert exporter.validate_input('/nonexistent/file.pdf') is False

    def test_validate_non_pdf(self, exporter, temp_dir):
        """Test validation of non-PDF file."""
        txt_path = Path(temp_dir) / 'test.txt'
        txt_path.write_text('not a pdf')
        assert exporter.validate_input(str(txt_path)) is False

    def test_validate_none_input(self, exporter):
        """Test validation of None input."""
        assert exporter.validate_input(None) is False

    def test_validate_empty_dict(self, exporter):
        """Test validation of empty dict."""
        assert exporter.validate_input({}) is False


# ===== Basic Conversion Tests =====

class TestBasicConversion:
    """Tests for basic PDF to Word conversion."""

    @pytest.mark.asyncio
    async def test_convert_simple_pdf(self, exporter, temp_dir):
        """Test converting a simple PDF to Word."""
        pdf_path = create_test_pdf(Path(temp_dir) / 'test.pdf')

        result = await exporter.process(str(pdf_path))

        assert result.success is True
        assert 'output_path' in result.data
        assert Path(result.data['output_path']).exists()
        assert result.data['output_path'].endswith('.docx')

    @pytest.mark.asyncio
    async def test_convert_with_dict_input(self, exporter, temp_dir):
        """Test converting with dict input."""
        pdf_path = create_test_pdf(Path(temp_dir) / 'test.pdf')
        input_data = {'pdf_path': str(pdf_path)}

        result = await exporter.process(input_data)

        assert result.success is True
        assert Path(result.data['output_path']).exists()

    @pytest.mark.asyncio
    async def test_convert_with_custom_output_name(self, exporter, temp_dir):
        """Test converting with custom output filename."""
        pdf_path = create_test_pdf(Path(temp_dir) / 'test.pdf')
        input_data = {
            'pdf_path': str(pdf_path),
            'output_name': 'custom_report.docx'
        }

        result = await exporter.process(input_data)

        assert result.success is True
        assert 'custom_report.docx' in result.data['output_path']

    @pytest.mark.asyncio
    async def test_convert_with_project_id(self, exporter, temp_dir):
        """Test converting with project ID."""
        pdf_path = create_test_pdf(Path(temp_dir) / 'test.pdf')
        input_data = {
            'pdf_path': str(pdf_path),
            'project_id': 'ODIC-2024-001'
        }

        result = await exporter.process(input_data)

        assert result.success is True
        assert 'ODIC-2024-001' in result.data['output_path']

    @pytest.mark.asyncio
    async def test_output_is_valid_docx(self, exporter, temp_dir):
        """Test that output is a valid Word document."""
        pdf_path = create_test_pdf(Path(temp_dir) / 'test.pdf')

        result = await exporter.process(str(pdf_path))

        assert result.success is True

        # Try to open as Word document
        doc = Document(result.data['output_path'])
        assert doc is not None


# ===== Multi-page PDF Tests =====

class TestMultiPageConversion:
    """Tests for multi-page PDF conversion."""

    @pytest.mark.asyncio
    async def test_convert_multipage_pdf(self, exporter, temp_dir):
        """Test converting a multi-page PDF."""
        pdf_path = create_test_pdf(Path(temp_dir) / 'multipage.pdf', num_pages=5)

        result = await exporter.process(str(pdf_path))

        assert result.success is True
        assert result.data.get('page_count', 0) == 5 or result.data.get('page_count', 0) > 0

    @pytest.mark.asyncio
    async def test_convert_single_page_pdf(self, exporter, temp_dir):
        """Test converting a single-page PDF."""
        pdf_path = create_test_pdf(Path(temp_dir) / 'single.pdf', num_pages=1)

        result = await exporter.process(str(pdf_path))

        assert result.success is True


# ===== Metadata Tests =====

class TestMetadata:
    """Tests for conversion metadata."""

    @pytest.mark.asyncio
    async def test_returns_original_pdf_path(self, exporter, temp_dir):
        """Test that result includes original PDF path."""
        pdf_path = create_test_pdf(Path(temp_dir) / 'test.pdf')

        result = await exporter.process(str(pdf_path))

        assert result.success is True
        assert result.data['original_pdf'] == str(pdf_path)

    @pytest.mark.asyncio
    async def test_returns_file_sizes(self, exporter, temp_dir):
        """Test that result includes file sizes."""
        pdf_path = create_test_pdf(Path(temp_dir) / 'test.pdf')

        result = await exporter.process(str(pdf_path))

        assert result.success is True
        assert 'original_size' in result.data
        assert 'output_size' in result.data
        assert result.data['original_size'] > 0
        assert result.data['output_size'] > 0

    @pytest.mark.asyncio
    async def test_returns_conversion_method(self, exporter, temp_dir):
        """Test that result includes conversion method."""
        pdf_path = create_test_pdf(Path(temp_dir) / 'test.pdf')

        result = await exporter.process(str(pdf_path))

        assert result.success is True
        assert 'conversion_method' in result.data


# ===== Synchronous Interface Tests =====

class TestSyncInterface:
    """Tests for synchronous export interface."""

    def test_sync_export(self, exporter, temp_dir):
        """Test synchronous export method."""
        pdf_path = create_test_pdf(Path(temp_dir) / 'test.pdf')

        success, output_path, metadata = exporter.export_sync(str(pdf_path))

        assert success is True
        assert output_path is not None
        assert Path(output_path).exists()

    def test_sync_export_with_project_id(self, exporter, temp_dir):
        """Test synchronous export with project ID."""
        pdf_path = create_test_pdf(Path(temp_dir) / 'test.pdf')

        success, output_path, metadata = exporter.export_sync(
            str(pdf_path),
            project_id='TEST-001'
        )

        assert success is True
        assert 'TEST-001' in output_path

    def test_sync_export_failure(self, exporter):
        """Test synchronous export with invalid input."""
        success, output_path, metadata = exporter.export_sync('/nonexistent.pdf')

        assert success is False
        assert output_path is None
        assert 'error' in metadata


# ===== Batch Export Tests =====

class TestBatchExport:
    """Tests for batch export functionality."""

    @pytest.mark.asyncio
    async def test_batch_export_multiple_pdfs(self, exporter, temp_dir):
        """Test batch export of multiple PDFs."""
        pdf1 = create_test_pdf(Path(temp_dir) / 'test1.pdf')
        pdf2 = create_test_pdf(Path(temp_dir) / 'test2.pdf')
        pdf3 = create_test_pdf(Path(temp_dir) / 'test3.pdf')

        result = await exporter.batch_export([str(pdf1), str(pdf2), str(pdf3)])

        assert result.success is True
        assert result.data['total'] == 3
        assert result.data['successful'] == 3
        assert result.data['failed'] == 0
        assert len(result.data['results']) == 3

    @pytest.mark.asyncio
    async def test_batch_export_with_project_id(self, exporter, temp_dir):
        """Test batch export with project ID."""
        pdf1 = create_test_pdf(Path(temp_dir) / 'test1.pdf')
        pdf2 = create_test_pdf(Path(temp_dir) / 'test2.pdf')

        result = await exporter.batch_export(
            [str(pdf1), str(pdf2)],
            project_id='BATCH-001'
        )

        assert result.success is True
        # All outputs should contain project ID
        for item in result.data['results']:
            if item['success']:
                assert 'BATCH-001' in item['output']

    @pytest.mark.asyncio
    async def test_batch_export_partial_failure(self, exporter, temp_dir):
        """Test batch export with some failures."""
        pdf1 = create_test_pdf(Path(temp_dir) / 'test1.pdf')
        bad_path = '/nonexistent/file.pdf'

        result = await exporter.batch_export([str(pdf1), bad_path])

        assert result.success is True  # Batch still succeeds
        assert result.data['successful'] >= 1
        assert result.data['failed'] >= 1


# ===== Export Info Tests =====

class TestExportInfo:
    """Tests for get_export_info method."""

    @pytest.mark.asyncio
    async def test_get_export_info(self, exporter, temp_dir):
        """Test getting info about exported Word document."""
        pdf_path = create_test_pdf(Path(temp_dir) / 'test.pdf')

        result = await exporter.process(str(pdf_path))
        assert result.success is True

        info = exporter.get_export_info(result.data['output_path'])

        assert info is not None
        assert 'path' in info
        assert 'size' in info
        assert 'modified' in info
        assert 'paragraph_count' in info

    def test_get_export_info_nonexistent(self, exporter):
        """Test getting info for nonexistent file."""
        info = exporter.get_export_info('/nonexistent/file.docx')
        assert info is None


# ===== Error Handling Tests =====

class TestErrorHandling:
    """Tests for error handling."""

    @pytest.mark.asyncio
    async def test_invalid_pdf_path(self, exporter):
        """Test handling of invalid PDF path."""
        result = await exporter.process('/nonexistent/file.pdf')

        assert result.success is False
        assert result.error is not None

    @pytest.mark.asyncio
    async def test_corrupted_pdf(self, exporter, temp_dir):
        """Test handling of corrupted PDF file."""
        corrupted_pdf = Path(temp_dir) / 'corrupted.pdf'
        corrupted_pdf.write_bytes(b'not a real pdf content')

        result = await exporter.process(str(corrupted_pdf))

        # Should fail gracefully
        assert result.success is False or 'warnings' in result.data


# ===== Fallback Mode Tests =====

class TestFallbackMode:
    """Tests for fallback conversion mode (when pdf2docx unavailable)."""

    @pytest.mark.asyncio
    async def test_fallback_conversion(self, config, temp_dir):
        """Test fallback conversion using text extraction."""
        from skills.word_exporter import WordExporter

        exporter = WordExporter(config)
        # Force fallback mode
        exporter._pdf2docx_available = False

        pdf_path = create_test_pdf(Path(temp_dir) / 'test.pdf')

        result = await exporter.process(str(pdf_path))

        # Should still succeed with fallback
        assert result.success is True
        assert 'output_path' in result.data
        assert Path(result.data['output_path']).exists()

        # Should indicate fallback method
        if 'warnings' in result.data:
            assert any('fallback' in w.lower() for w in result.data['warnings'])


# ===== Output Path Uniqueness =====

class TestOutputPathUniqueness:
    """Tests for unique output paths."""

    @pytest.mark.asyncio
    async def test_output_path_uniqueness(self, exporter, temp_dir):
        """Test that multiple exports create unique output paths."""
        pdf_path = create_test_pdf(Path(temp_dir) / 'test.pdf')

        result1 = await exporter.process(str(pdf_path))
        result2 = await exporter.process(str(pdf_path))

        assert result1.success is True
        assert result2.success is True
        # Output paths should be different
        assert result1.data['output_path'] != result2.data['output_path']


# ===== Integration Tests =====

class TestIntegration:
    """Integration tests for full export workflow."""

    @pytest.mark.asyncio
    async def test_full_export_workflow(self, exporter, temp_dir):
        """Test complete export workflow."""
        # Create test PDF
        pdf_path = create_test_pdf(Path(temp_dir) / 'report.pdf', num_pages=5)

        # Export to Word
        result = await exporter.process({
            'pdf_path': str(pdf_path),
            'project_id': 'WORKFLOW-001',
            'output_name': 'final_report'
        })

        assert result.success is True
        output_path = result.data['output_path']

        # Verify output
        assert Path(output_path).exists()
        doc = Document(output_path)
        assert len(doc.paragraphs) > 0

        # Get export info
        info = exporter.get_export_info(output_path)
        assert info is not None
        assert info['paragraph_count'] > 0
