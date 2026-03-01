"""
ODIC ESA Pipeline - Word Exporter Skill

Converts assembled PDF reports back to Word documents (.docx) for manual editing.
Uses pdf2docx for PDF to Word conversion with style preservation.

This skill allows Rose to make manual edits to the assembled report before
final export to clients.
"""

import os
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from .base import BaseSkill, SkillResult


class WordExporter(BaseSkill):
    """
    Converts PDF reports to Word documents for manual editing.

    Uses pdf2docx for conversion, preserving:
    - Text content and formatting
    - Images and figures
    - Tables
    - Page layout (approximate)

    Output is a .docx file ready for editing in Microsoft Word.
    """

    def __init__(self, config: dict):
        """
        Initialize the Word exporter.

        Args:
            config: Configuration dictionary with optional 'word_exporter' section
        """
        super().__init__(config)

        exporter_config = config.get('word_exporter', {})

        # Output directory for exported Word files
        self.output_dir = Path(exporter_config.get(
            'output_dir',
            config.get('pipeline', {}).get('output_dir', './exports')
        ))
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # Export settings
        self.preserve_images = exporter_config.get('preserve_images', True)
        self.preserve_tables = exporter_config.get('preserve_tables', True)
        self.page_margin = exporter_config.get('page_margin_inches', 1.0)

        # Check if pdf2docx is available (optional dependency)
        self._pdf2docx_available = self._check_pdf2docx()

    def _check_pdf2docx(self) -> bool:
        """Check if pdf2docx library is available."""
        try:
            from pdf2docx import Converter
            return True
        except ImportError:
            self.logger.warning(
                "pdf2docx not installed. Install with: pip install pdf2docx"
            )
            return False

    def validate_input(self, input_data: Any) -> bool:
        """
        Validate that input is a valid PDF file path.

        Args:
            input_data: Path to PDF file (str or Path) or dict with 'pdf_path'

        Returns:
            True if input is valid
        """
        if isinstance(input_data, dict):
            pdf_path = input_data.get('pdf_path') or input_data.get('path')
        else:
            pdf_path = input_data

        if pdf_path is None:
            self.logger.error("No PDF path provided")
            return False

        path = Path(pdf_path)

        if not path.exists():
            self.logger.error(f"PDF file does not exist: {path}")
            return False

        if path.suffix.lower() != '.pdf':
            self.logger.error(f"File is not a PDF: {path}")
            return False

        return True

    async def process(self, input_data: Any) -> SkillResult:
        """
        Convert a PDF report to Word document.

        Args:
            input_data: Path to PDF file or dict with:
                - pdf_path: Path to the PDF file
                - output_name: Optional custom output filename
                - project_id: Optional project ID for naming

        Returns:
            SkillResult with output Word file path and metadata
        """
        # Parse input
        if isinstance(input_data, dict):
            pdf_path = Path(input_data.get('pdf_path') or input_data.get('path'))
            output_name = input_data.get('output_name')
            project_id = input_data.get('project_id')
        else:
            pdf_path = Path(input_data)
            output_name = None
            project_id = None

        if not self.validate_input(input_data):
            return SkillResult.fail(
                error="Invalid input - must provide valid PDF file path",
                data={'input': str(input_data)}
            )

        # Generate output filename
        if output_name:
            output_filename = output_name if output_name.endswith('.docx') else f"{output_name}.docx"
        elif project_id:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_filename = f"{project_id}_report_{timestamp}.docx"
        else:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_%f')
            output_filename = f"{pdf_path.stem}_{timestamp}.docx"

        output_path = self.output_dir / output_filename

        try:
            # Convert PDF to Word
            if self._pdf2docx_available:
                result = await self._convert_with_pdf2docx(pdf_path, output_path)
            else:
                result = await self._convert_with_fallback(pdf_path, output_path)

            if not result['success']:
                return SkillResult.fail(
                    error=result.get('error', 'Conversion failed'),
                    data={'pdf_path': str(pdf_path)}
                )

            # Verify output
            if not output_path.exists():
                return SkillResult.fail(
                    error="Conversion failed - output file not created",
                    data={'pdf_path': str(pdf_path)}
                )

            self.logger.info(
                f"Exported PDF to Word: {pdf_path.name} -> {output_path.name}"
            )

            return SkillResult.ok(
                data={
                    'output_path': str(output_path),
                    'original_pdf': str(pdf_path),
                    'original_size': pdf_path.stat().st_size,
                    'output_size': output_path.stat().st_size,
                    'conversion_method': result.get('method', 'unknown'),
                    'page_count': result.get('page_count', 0),
                    'warnings': result.get('warnings', [])
                }
            )

        except Exception as e:
            self.logger.exception(f"Word export failed for {pdf_path.name}")
            return SkillResult.fail(
                error=f"Export failed: {str(e)}",
                data={'pdf_path': str(pdf_path)}
            )

    async def _convert_with_pdf2docx(
        self,
        pdf_path: Path,
        output_path: Path
    ) -> Dict[str, Any]:
        """
        Convert PDF to Word using pdf2docx library.

        This provides high-fidelity conversion preserving:
        - Text formatting (fonts, sizes, colors)
        - Images and graphics
        - Tables
        - Page layout

        Args:
            pdf_path: Input PDF file path
            output_path: Output Word file path

        Returns:
            Dict with success status and metadata
        """
        from pdf2docx import Converter

        warnings = []
        page_count = 0

        try:
            # Create converter
            cv = Converter(str(pdf_path))

            # Get page count
            page_count = len(cv.pages)

            # Convert all pages
            cv.convert(str(output_path), start=0, end=None)
            cv.close()

            return {
                'success': True,
                'method': 'pdf2docx',
                'page_count': page_count,
                'warnings': warnings
            }

        except Exception as e:
            self.logger.error(f"pdf2docx conversion failed: {e}")
            return {
                'success': False,
                'error': str(e),
                'method': 'pdf2docx'
            }

    async def _convert_with_fallback(
        self,
        pdf_path: Path,
        output_path: Path
    ) -> Dict[str, Any]:
        """
        Fallback conversion when pdf2docx is not available.

        Creates a simple Word document with extracted text content.
        Does not preserve complex formatting but maintains content.

        Args:
            pdf_path: Input PDF file path
            output_path: Output Word file path

        Returns:
            Dict with success status and metadata
        """
        warnings = ["Using fallback conversion - formatting may be limited"]

        try:
            # Try to use pdfplumber for text extraction
            try:
                import pdfplumber
                use_pdfplumber = True
            except ImportError:
                use_pdfplumber = False
                warnings.append("pdfplumber not available - using PyPDF2")

            # Extract text from PDF
            text_content = []
            page_count = 0

            if use_pdfplumber:
                with pdfplumber.open(str(pdf_path)) as pdf:
                    page_count = len(pdf.pages)
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_content.append(text)
            else:
                # Fallback to PyPDF2
                try:
                    from PyPDF2 import PdfReader
                    reader = PdfReader(str(pdf_path))
                    page_count = len(reader.pages)
                    for page in reader.pages:
                        text = page.extract_text()
                        if text:
                            text_content.append(text)
                except ImportError:
                    return {
                        'success': False,
                        'error': "No PDF text extraction library available",
                        'method': 'fallback'
                    }

            # Create Word document
            doc = Document()

            # Set up styles
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

            # Add content
            for i, page_text in enumerate(text_content):
                if i > 0:
                    # Add page break between pages
                    doc.add_page_break()

                # Split into paragraphs and add
                paragraphs = page_text.split('\n\n')
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:
                        # Check if it looks like a heading
                        if len(para_text) < 100 and para_text.isupper():
                            p = doc.add_heading(para_text, level=1)
                        elif para_text.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                            # Section heading
                            p = doc.add_heading(para_text, level=2)
                        else:
                            p = doc.add_paragraph(para_text)

            # Save document
            doc.save(str(output_path))

            warnings.append("Complex formatting (images, tables) not preserved in fallback mode")

            return {
                'success': True,
                'method': 'fallback_text_extraction',
                'page_count': page_count,
                'warnings': warnings
            }

        except Exception as e:
            self.logger.error(f"Fallback conversion failed: {e}")
            return {
                'success': False,
                'error': str(e),
                'method': 'fallback'
            }

    def export_sync(
        self,
        pdf_path: Union[str, Path],
        output_name: Optional[str] = None,
        project_id: Optional[str] = None
    ) -> Tuple[bool, Optional[str], Dict[str, Any]]:
        """
        Synchronous export method for simpler use cases.

        Args:
            pdf_path: Path to PDF file
            output_name: Optional output filename
            project_id: Optional project ID for naming

        Returns:
            Tuple of (success, output_path, metadata)
        """
        import asyncio

        input_data = {
            'pdf_path': str(pdf_path),
            'output_name': output_name,
            'project_id': project_id
        }

        result = asyncio.run(self.process(input_data))

        if result.success:
            return True, result.data.get('output_path'), result.data
        else:
            return False, None, {'error': result.error}

    async def batch_export(
        self,
        pdf_paths: List[Union[str, Path]],
        project_id: Optional[str] = None
    ) -> SkillResult:
        """
        Export multiple PDFs to Word documents.

        Args:
            pdf_paths: List of PDF file paths
            project_id: Optional project ID for naming

        Returns:
            SkillResult with list of exported files
        """
        results = []
        successful = 0
        failed = 0

        for pdf_path in pdf_paths:
            input_data = {
                'pdf_path': str(pdf_path),
                'project_id': project_id
            }

            result = await self.process(input_data)
            results.append({
                'input': str(pdf_path),
                'success': result.success,
                'output': result.data.get('output_path') if result.success else None,
                'error': result.error
            })

            if result.success:
                successful += 1
            else:
                failed += 1

        return SkillResult.ok(
            data={
                'total': len(pdf_paths),
                'successful': successful,
                'failed': failed,
                'results': results
            }
        )

    def get_export_info(self, output_path: Union[str, Path]) -> Optional[Dict[str, Any]]:
        """
        Get information about an exported Word document.

        Args:
            output_path: Path to the exported Word file

        Returns:
            Dict with file info or None if not found
        """
        path = Path(output_path)

        if not path.exists():
            return None

        try:
            doc = Document(str(path))
            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)

            # Count images (inline shapes)
            image_count = 0
            for para in doc.paragraphs:
                for run in para.runs:
                    if run._element.xpath('.//a:blip'):
                        image_count += 1

            return {
                'path': str(path),
                'size': path.stat().st_size,
                'modified': datetime.fromtimestamp(path.stat().st_mtime).isoformat(),
                'paragraph_count': paragraph_count,
                'table_count': table_count,
                'image_count': image_count
            }

        except Exception as e:
            self.logger.error(f"Failed to get export info: {e}")
            return None
