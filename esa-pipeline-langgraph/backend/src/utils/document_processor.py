"""
Document Processing Utilities

Handles:
- PDF text extraction (PyMuPDF)
- Word document handling (python-docx)
- OCR for image-based documents (pytesseract)
- File format normalization
"""

import os
import io
import hashlib
import logging
from pathlib import Path
from typing import Tuple, Optional, Dict, Any, List
from dataclasses import dataclass

logger = logging.getLogger(__name__)

# Try to import optional dependencies
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logger.warning("PyMuPDF not available - PDF processing limited")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available - Word document processing limited")

try:
    from PIL import Image
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    logger.warning("pytesseract/PIL not available - OCR not available")


@dataclass
class ProcessedDocument:
    """Result of processing a document."""
    text_content: str
    page_count: int
    format: str
    ocr_confidence: Optional[float] = None
    metadata: Dict[str, Any] = None
    images: List[bytes] = None  # Extracted images if any


def compute_file_hash(file_path: str) -> str:
    """Compute SHA256 hash of file for deduplication."""
    sha256_hash = hashlib.sha256()
    with open(file_path, "rb") as f:
        for byte_block in iter(lambda: f.read(4096), b""):
            sha256_hash.update(byte_block)
    return sha256_hash.hexdigest()


def get_file_format(file_path: str) -> str:
    """Determine file format from extension and magic bytes."""
    path = Path(file_path)
    ext = path.suffix.lower()

    format_map = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'doc',
        '.jpg': 'image',
        '.jpeg': 'image',
        '.png': 'image',
        '.tiff': 'image',
        '.tif': 'image',
        '.bmp': 'image',
    }

    return format_map.get(ext, 'unknown')


def process_document(file_path: str) -> ProcessedDocument:
    """
    Process any supported document format.

    Extracts text content, page count, and metadata.
    Uses OCR for image-based documents or scanned PDFs.
    """
    format = get_file_format(file_path)
    file_size = os.path.getsize(file_path)

    if format == 'pdf':
        return process_pdf(file_path)
    elif format == 'docx':
        return process_docx(file_path)
    elif format == 'doc':
        return process_legacy_doc(file_path)
    elif format == 'image':
        return process_image(file_path)
    else:
        logger.warning(f"Unsupported format: {format} for {file_path}")
        return ProcessedDocument(
            text_content="",
            page_count=0,
            format=format,
            metadata={"error": "Unsupported format"}
        )


def process_pdf(file_path: str) -> ProcessedDocument:
    """Process PDF document."""
    if not PYMUPDF_AVAILABLE:
        return ProcessedDocument(
            text_content="",
            page_count=0,
            format='pdf',
            metadata={"error": "PyMuPDF not available"}
        )

    try:
        doc = fitz.open(file_path)
        text_parts = []
        page_count = len(doc)
        ocr_confidence = None
        images = []

        for page_num, page in enumerate(doc):
            # Try to extract text
            text = page.get_text()

            if text.strip():
                text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
            else:
                # Page might be scanned - try OCR
                if TESSERACT_AVAILABLE:
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    ocr_text, conf = ocr_image_bytes(img_data)
                    if ocr_text:
                        text_parts.append(f"--- Page {page_num + 1} (OCR) ---\n{ocr_text}")
                        if ocr_confidence is None:
                            ocr_confidence = conf
                        else:
                            ocr_confidence = (ocr_confidence + conf) / 2

            # Extract embedded images if any
            image_list = page.get_images()
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    images.append(base_image["image"])
                except Exception:
                    pass

        doc.close()

        metadata = {
            "pages": page_count,
            "has_text_layer": any("(OCR)" not in p for p in text_parts),
            "image_count": len(images),
        }

        return ProcessedDocument(
            text_content="\n\n".join(text_parts),
            page_count=page_count,
            format='pdf',
            ocr_confidence=ocr_confidence,
            metadata=metadata,
            images=images,
        )

    except Exception as e:
        logger.error(f"Failed to process PDF {file_path}: {e}")
        return ProcessedDocument(
            text_content="",
            page_count=0,
            format='pdf',
            metadata={"error": str(e)}
        )


def process_docx(file_path: str) -> ProcessedDocument:
    """Process Word document."""
    if not DOCX_AVAILABLE:
        return ProcessedDocument(
            text_content="",
            page_count=0,
            format='docx',
            metadata={"error": "python-docx not available"}
        )

    try:
        doc = Document(file_path)
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                # Check if it's a heading
                if para.style.name.startswith('Heading'):
                    text_parts.append(f"\n## {para.text}\n")
                else:
                    text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                table_text.append(" | ".join(row_text))
            if table_text:
                text_parts.append("\n[TABLE]\n" + "\n".join(table_text) + "\n[/TABLE]\n")

        # Estimate page count (rough - Word doesn't store this reliably)
        total_chars = sum(len(p) for p in text_parts)
        estimated_pages = max(1, total_chars // 3000)  # ~3000 chars per page

        metadata = {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "sections": len(doc.sections),
        }

        return ProcessedDocument(
            text_content="\n".join(text_parts),
            page_count=estimated_pages,
            format='docx',
            metadata=metadata,
        )

    except Exception as e:
        logger.error(f"Failed to process DOCX {file_path}: {e}")
        return ProcessedDocument(
            text_content="",
            page_count=0,
            format='docx',
            metadata={"error": str(e)}
        )


def process_legacy_doc(file_path: str) -> ProcessedDocument:
    """Process legacy .doc file."""
    # Legacy .doc files are harder to process without antiword or similar
    logger.warning(f"Legacy .doc format not fully supported: {file_path}")
    return ProcessedDocument(
        text_content="[Legacy .doc format - please convert to .docx]",
        page_count=0,
        format='doc',
        metadata={"error": "Legacy format not supported"}
    )


def process_image(file_path: str) -> ProcessedDocument:
    """Process image file with OCR."""
    if not TESSERACT_AVAILABLE:
        return ProcessedDocument(
            text_content="",
            page_count=1,
            format='image',
            metadata={"error": "Tesseract OCR not available"}
        )

    try:
        with open(file_path, 'rb') as f:
            image_data = f.read()

        text, confidence = ocr_image_bytes(image_data)

        return ProcessedDocument(
            text_content=text,
            page_count=1,
            format='image',
            ocr_confidence=confidence,
            metadata={"ocr_applied": True},
            images=[image_data],
        )

    except Exception as e:
        logger.error(f"Failed to process image {file_path}: {e}")
        return ProcessedDocument(
            text_content="",
            page_count=1,
            format='image',
            metadata={"error": str(e)}
        )


def ocr_image_bytes(image_data: bytes) -> Tuple[str, float]:
    """Perform OCR on image bytes."""
    if not TESSERACT_AVAILABLE:
        return "", 0.0

    try:
        image = Image.open(io.BytesIO(image_data))

        # Get OCR data with confidence scores
        ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)

        # Extract text and calculate average confidence
        texts = []
        confidences = []

        for i, text in enumerate(ocr_data['text']):
            if text.strip():
                texts.append(text)
                conf = ocr_data['conf'][i]
                if conf > 0:  # -1 means no confidence available
                    confidences.append(conf)

        full_text = " ".join(texts)
        avg_confidence = sum(confidences) / len(confidences) / 100 if confidences else 0.5

        return full_text, avg_confidence

    except Exception as e:
        logger.error(f"OCR failed: {e}")
        return "", 0.0


def merge_pdfs(input_paths: List[str], output_path: str) -> int:
    """
    Merge multiple PDFs into one.

    Returns total page count of merged document.
    """
    if not PYMUPDF_AVAILABLE:
        raise RuntimeError("PyMuPDF required for PDF merging")

    try:
        merged_doc = fitz.open()
        total_pages = 0

        for path in input_paths:
            if not os.path.exists(path):
                logger.warning(f"Skipping missing file: {path}")
                continue

            try:
                doc = fitz.open(path)
                merged_doc.insert_pdf(doc)
                total_pages += len(doc)
                doc.close()
            except Exception as e:
                logger.error(f"Failed to merge {path}: {e}")

        merged_doc.save(output_path)
        merged_doc.close()

        logger.info(f"Merged {len(input_paths)} PDFs into {output_path} ({total_pages} pages)")
        return total_pages

    except Exception as e:
        logger.error(f"PDF merge failed: {e}")
        raise


def split_pdf(input_path: str, output_dir: str, split_points: List[int]) -> List[str]:
    """
    Split a PDF at specified page numbers.

    Args:
        input_path: Path to input PDF
        output_dir: Directory for output files
        split_points: List of page numbers to split at (1-indexed)

    Returns:
        List of output file paths
    """
    if not PYMUPDF_AVAILABLE:
        raise RuntimeError("PyMuPDF required for PDF splitting")

    try:
        doc = fitz.open(input_path)
        total_pages = len(doc)

        # Normalize split points
        split_points = sorted(set([0] + split_points + [total_pages]))

        output_paths = []
        base_name = Path(input_path).stem

        for i in range(len(split_points) - 1):
            start_page = split_points[i]
            end_page = split_points[i + 1]

            part_doc = fitz.open()
            part_doc.insert_pdf(doc, from_page=start_page, to_page=end_page - 1)

            part_num = i + 1
            total_parts = len(split_points) - 1
            output_name = f"{base_name}_Part{part_num}of{total_parts}.pdf"
            output_path = os.path.join(output_dir, output_name)

            part_doc.save(output_path)
            part_doc.close()
            output_paths.append(output_path)

        doc.close()

        logger.info(f"Split {input_path} into {len(output_paths)} parts")
        return output_paths

    except Exception as e:
        logger.error(f"PDF split failed: {e}")
        raise


def get_pdf_page_count(file_path: str) -> int:
    """Get page count of a PDF without loading all content."""
    if not PYMUPDF_AVAILABLE:
        return 0

    try:
        doc = fitz.open(file_path)
        count = len(doc)
        doc.close()
        return count
    except Exception:
        return 0


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract all text content from a PDF file.

    Args:
        file_path: Path to the PDF file

    Returns:
        Extracted text content as a string
    """
    if not PYMUPDF_AVAILABLE:
        logger.warning("PyMuPDF not available for text extraction")
        return ""

    try:
        doc = fitz.open(file_path)
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"Failed to extract text from {file_path}: {e}")
        return ""


def extract_pdf_pages(input_path: str, page_range: Tuple[int, int], output_path: str) -> bool:
    """
    Extract a range of pages from a PDF.

    Args:
        input_path: Source PDF path
        page_range: (start, end) 0-indexed, end exclusive
        output_path: Where to save extracted pages

    Returns:
        True if successful
    """
    if not PYMUPDF_AVAILABLE:
        return False

    try:
        doc = fitz.open(input_path)
        new_doc = fitz.open()
        new_doc.insert_pdf(doc, from_page=page_range[0], to_page=page_range[1] - 1)
        new_doc.save(output_path)
        new_doc.close()
        doc.close()
        return True
    except Exception as e:
        logger.error(f"Page extraction failed: {e}")
        return False
