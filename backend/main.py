"""Environmental Report Assembler - FastAPI Backend."""

from __future__ import annotations

import asyncio
import json
import logging
import os
import re
import shutil
import uuid
import zipfile
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Optional

import httpx

from fastapi import (
    Depends,
    FastAPI,
    File,
    Form,
    HTTPException,
    Query,
    UploadFile,
)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.orm import Session
from sse_starlette.sse import EventSourceResponse

from assembler import assemble_report
from classifier import classify_document, classify_by_filename_legacy, classify_all_documents_queued, validate_assembly
from report_director import run_report_director
from compressor import compress_pdf, get_file_size_display
from config import settings
from converter import convert_to_pdf, get_pdf_page_count, async_convert_to_pdf, async_get_pdf_page_count
from database import Document, Report, ChatMessage, ActionSnapshot, SessionLocal, get_db, init_db
from models import (
    AssembleRequest,
    BatchUpdateRequest,
    ChatRequest,
    CompressRequest,
    DeletePagesRequest,
    DocxContentResponse,
    DocxContentUpdateRequest,
    DocxParagraph,
    DocxRun,
    DocumentResponse,
    DocumentStatus,
    DocumentUpdate,
    ReorderRequest,
    ReportCreate,
    ReportResponse,
    ReportStatus,
    ReportUpdate,
    SectionCategory,
    TextReplaceRequest,
)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# ---- Bug 4: Compiled report detection via content fingerprinting ----

def is_compiled_report(pdf_path: Path) -> bool:
    """Detect if a PDF is a previously-assembled Phase I ESA report by checking
    for Phase I ESA-specific structural markers in the first 30 pages.
    Uses ESA-specific appendix labels (not generic 'appendix a/b/c') to avoid
    false positives on other environmental reports that have their own appendices.
    Does NOT use page count as a heuristic."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(pdf_path))
        text_parts = []
        for page in reader.pages[:30]:
            text = page.extract_text() or ""
            text_parts.append(text)
        full_text = "\n".join(text_parts).lower()

        # Phase I ESA-specific appendix content labels (not generic a/b/c/d)
        esa_section_markers = [
            ("property location map", "plot plan"),     # Appendix A
            ("site photographs", "photographs"),         # Appendix B
            ("database report", "radius report"),        # Appendix C
            ("historical records",),                     # Appendix D
            ("agency records", "public agency"),         # Appendix E
            ("qualifications of environmental",),        # Appendix F
        ]
        esa_marker_count = sum(
            1 for markers in esa_section_markers
            if any(m in full_text for m in markers)
        )

        has_toc = "table of contents" in full_text
        has_phase_i = "phase i" in full_text or "phase 1 environmental" in full_text

        # Require 3+ ESA-specific section labels AND TOC AND Phase I reference
        return esa_marker_count >= 3 and has_toc and has_phase_i

    except Exception as e:
        logger.warning(f"Compiled report check failed for {pdf_path.name}: {e}")
        return False


# ---- Bug 5: Document version deduplication ----

def _normalize_basename(filename: str) -> str:
    """Strip explicit revision indicators from a filename to get a normalized base name.
    Only strips clear version markers (v2, rev1, FINAL, draft, copy, (1), (Revised)).
    Does NOT strip trailing numbers that could be dates, photo numbers, well IDs, etc.
    Returns without extension so .docx and .pdf versions of the same doc group together."""
    name = Path(filename).stem.lower()
    # Strip explicit version/revision markers: _v2, -rev1, _version3
    name = re.sub(r'[_\-\s]*(?:v|rev|version)\s*\d+', '', name)
    # Strip parenthesized version markers: (1), (2), (Revised), (Updated), (Final)
    name = re.sub(r'\s*\((?:revised|updated|final|draft|copy|\d+)\s*\d*\)', '', name, flags=re.IGNORECASE)
    # Strip trailing keywords: _FINAL, _draft, _copy, _revised, _updated
    name = re.sub(r'[_\-\s]*(?:final|draft|copy|revised|updated)$', '', name, flags=re.IGNORECASE)
    # Strip reviewer initials pattern: -rev-mam, _rev_jdm
    name = re.sub(r'[_\-]rev[_\-][a-z]{2,4}$', '', name)
    # Strip trailing separators
    name = re.sub(r'[_\-\s]+$', '', name)
    return name


def deduplicate_documents(report_id: int, db: Session):
    """Group documents by normalized base name. For groups with multiple files,
    keep only the newest (by file modified time / upload timestamp) and set
    is_included=False on older versions."""
    docs = db.query(Document).filter(
        Document.report_id == report_id,
        Document.is_included == True,
    ).all()

    # Group by normalized name
    groups: dict[str, list[Document]] = {}
    for doc in docs:
        key = _normalize_basename(doc.original_filename)
        groups.setdefault(key, []).append(doc)

    deduped_count = 0
    for key, group in groups.items():
        if len(group) <= 1:
            continue

        # Sort by original file mtime (preserved by shutil.copy2), fallback to created_at
        def _file_mtime(doc):
            stored = settings.UPLOAD_DIR / str(doc.report_id) / "originals" / doc.stored_filename
            try:
                return stored.stat().st_mtime
            except Exception:
                return 0
        group.sort(key=lambda d: _file_mtime(d), reverse=True)
        newest = group[0]

        for older in group[1:]:
            older.is_included = False
            older.reasoning = f"Superseded by newer version: {newest.original_filename}"
            deduped_count += 1

    if deduped_count > 0:
        db.commit()
        logger.info(f"Deduplicated {deduped_count} older document versions for report {report_id}")


app = FastAPI(title="Environmental Report Assembler", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins since we're serving the frontend ourselves
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Serve static frontend files from React build
from fastapi.staticfiles import StaticFiles

# Static files are at: backend/static (copied from frontend/dist by Dockerfile)
# Using relative path so it works in all deployment environments
static_dir = Path(__file__).parent / "static"

if static_dir.exists() and any(static_dir.iterdir()):
    app.mount("/", StaticFiles(directory=str(static_dir), html=True), name="frontend")
    logger.info(f"✓ Mounted frontend static files from {static_dir} ({len(list(static_dir.glob('**/*')))} files)")
else:
    logger.warning(
        f"Static directory not fully initialized at {static_dir}. "
        "Frontend will not be served. Make sure Dockerfile.prod copies frontend/dist to backend/static. "
        "For API-only deployments, this warning is OK."
    )


# ---- Global Exception Handlers ----

from fastapi.responses import JSONResponse
from starlette.exceptions import HTTPException as StarletteHTTPException


@app.exception_handler(StarletteHTTPException)
async def http_exception_handler(request, exc):
    """Handle HTTP exceptions with structured JSON response."""
    return JSONResponse(
        status_code=exc.status_code,
        content={
            "error": exc.detail,
            "status_code": exc.status_code,
            "path": str(request.url.path),
        },
    )


@app.exception_handler(Exception)
async def general_exception_handler(request, exc):
    """Handle unexpected exceptions with useful error context."""
    logger.error(f"Unhandled exception: {exc}", exc_info=True)
    return JSONResponse(
        status_code=500,
        content={
            "error": "Internal server error",
            "detail": str(exc) if not isinstance(exc, HTTPException) else exc.detail,
            "type": type(exc).__name__,
            "path": str(request.url.path),
        },
    )


@app.on_event("startup")
def startup():
    """Startup event: Initialize database and validate configuration."""
    # Initialize database with error handling
    try:
        init_db()
        logger.info("✓ Database initialized successfully")
    except Exception as e:
        logger.error(f"✗ Failed to initialize database: {e}", exc_info=True)
        raise RuntimeError(f"Database initialization failed: {e}") from e
    
    # Validate configuration
    try:
        # If AI_BACKEND is anthropic, require API key
        if settings.AI_BACKEND == "anthropic":
            if not settings.ANTHROPIC_API_KEY:
                logger.warning(
                    "AI_BACKEND=anthropic but ANTHROPIC_API_KEY not set. "
                    "Falling back to ollama. Set ANTHROPIC_API_KEY environment variable to use Claude."
                )
                settings.AI_BACKEND = "ollama"
            else:
                logger.info("✓ Anthropic API key configured")
        
        # Validate upload directory exists
        settings.UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
        logger.info(f"✓ Upload directory ready: {settings.UPLOAD_DIR}")
        
        # Log active backend
        logger.info(f"✓ AI Backend: {settings.AI_BACKEND}")
        
    except Exception as e:
        logger.error(f"✗ Configuration validation failed: {e}", exc_info=True)
        raise RuntimeError(f"Configuration validation failed: {e}") from e


@app.get("/health")
def health_check():
    """Health check endpoint for Docker/Render deployment."""
    return {"status": "ok"}


# ---- Report CRUD ----

@app.post("/api/reports", response_model=ReportResponse)
def create_report(data: ReportCreate, db: Session = Depends(get_db)):
    report = Report(
        name=data.name,
        address=data.address,
        project_number=data.project_number,
        has_reliance_letter=data.has_reliance_letter,
    )
    db.add(report)
    db.commit()
    db.refresh(report)
    return _report_to_response(report)


@app.get("/api/reports", response_model=list[ReportResponse])
def list_reports(db: Session = Depends(get_db)):
    reports = db.query(Report).order_by(Report.created_at.desc()).all()
    return [_report_to_response(r) for r in reports]


@app.get("/api/reports/{report_id}", response_model=ReportResponse)
def get_report(report_id: int, db: Session = Depends(get_db)):
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "Report not found")
    return _report_to_response(report)


@app.put("/api/reports/{report_id}", response_model=ReportResponse)
def update_report(report_id: int, data: ReportUpdate, db: Session = Depends(get_db)):
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "Report not found")

    for field, value in data.dict(exclude_unset=True).items():
        if value is not None:
            setattr(report, field, value.value if hasattr(value, "value") else value)

    report.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(report)
    return _report_to_response(report)


@app.delete("/api/reports/{report_id}")
def delete_report(report_id: int, db: Session = Depends(get_db)):
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "Report not found")

    # Delete files
    report_dir = settings.UPLOAD_DIR / str(report_id)
    if report_dir.exists():
        shutil.rmtree(report_dir)

    db.delete(report)
    db.commit()
    return {"status": "deleted"}


# ---- Document Upload ----

@app.post("/api/reports/{report_id}/upload")
async def upload_files(
    report_id: int,
    files: list[UploadFile] = File(...),
    db: Session = Depends(get_db),
):
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "Report not found")

    report_dir = settings.UPLOAD_DIR / str(report_id) / "originals"
    report_dir.mkdir(parents=True, exist_ok=True)
    pdf_dir = settings.UPLOAD_DIR / str(report_id) / "pdfs"
    pdf_dir.mkdir(parents=True, exist_ok=True)

    uploaded = []

    # Collect all files to process (expanding zips inline)
    file_items: list[tuple[str, bytes, str]] = []  # (filename, content, rel_path)

    for file in files:
        content = await file.read()
        ext = Path(file.filename).suffix.lower()

        if ext == ".zip":
            # Extract zip and queue each inner file
            try:
                with zipfile.ZipFile(BytesIO(content)) as zf:
                    for info in zf.infolist():
                        if info.is_dir():
                            continue
                        inner_name = Path(info.filename).name
                        if not inner_name or inner_name.startswith("._") or "__MACOSX" in info.filename:
                            continue
                        inner_ext = Path(inner_name).suffix.lower()
                        if inner_ext in settings.SKIP_EXTENSIONS:
                            continue
                        file_items.append((inner_name, zf.read(info.filename), info.filename))
            except zipfile.BadZipFile:
                logger.warning(f"Bad zip file: {file.filename}")
            continue

        file_items.append((file.filename, content, ""))

    # ── Phase 1: Save files to disk + create DB records (fast) ──
    for filename, content, rel_path in file_items:
        ext = Path(filename).suffix.lower()

        if ext in settings.SKIP_EXTENSIONS:
            continue
        if ext not in settings.SUPPORTED_EXTENSIONS:
            continue

        # Deduplicate: if same filename already exists in this report, update it
        existing = db.query(Document).filter(
            Document.report_id == report_id,
            Document.original_filename == filename,
        ).first()

        if existing:
            stored_path = report_dir / existing.stored_filename
            stored_path.write_bytes(content)
            existing.file_size = len(content)
            existing.original_path = rel_path if rel_path else existing.original_path
            quick_class = classify_by_filename_legacy(filename, rel_path)
            existing.category = quick_class.category.value if quick_class else SectionCategory.UNCLASSIFIED.value
            existing.subcategory = quick_class.subcategory if quick_class else None
            existing.confidence = quick_class.confidence if quick_class else None
            existing.reasoning = quick_class.reasoning if quick_class else None
            existing.status = DocumentStatus.UPLOADED.value
            doc = existing
        else:
            # Generate unique stored filename
            stored_name = f"{uuid.uuid4().hex}{ext}"
            stored_path = report_dir / stored_name
            stored_path.write_bytes(content)

            # Quick filename-based classification
            quick_class = classify_by_filename_legacy(filename, rel_path)
            category = quick_class.category.value if quick_class else SectionCategory.UNCLASSIFIED.value
            subcategory = quick_class.subcategory if quick_class else None
            confidence = quick_class.confidence if quick_class else None

            doc = Document(
                report_id=report_id,
                original_filename=filename,
                original_path=rel_path if rel_path else None,
                stored_filename=stored_name,
                file_size=len(content),
                category=category,
                subcategory=subcategory,
                confidence=confidence,
                reasoning=quick_class.reasoning if quick_class else None,
                status=DocumentStatus.UPLOADED.value,
            )
            db.add(doc)

        uploaded.append(doc)

    # Single batch commit for all saved files
    db.commit()
    for doc in uploaded:
        db.refresh(doc)

    # Collect doc info needed by background task before response returns
    doc_ids = [doc.id for doc in uploaded]

    report.status = ReportStatus.IN_PROGRESS.value
    report.updated_at = datetime.utcnow()
    db.commit()

    # Return immediately — user sees documents in UPLOADED state
    response = {"uploaded": len(uploaded), "documents": [_doc_to_response(d) for d in uploaded]}

    # ── Phase 2: Background processing (conversion, page count, compiled check, dedup) ──
    asyncio.create_task(_process_uploaded_docs(report_id, doc_ids, report_dir, pdf_dir))

    return response


# Max concurrent conversions (LibreOffice is CPU-heavy)
_CONVERSION_SEMAPHORE: Optional[asyncio.Semaphore] = None


def _get_conversion_semaphore() -> asyncio.Semaphore:
    global _CONVERSION_SEMAPHORE
    if _CONVERSION_SEMAPHORE is None:
        _CONVERSION_SEMAPHORE = asyncio.Semaphore(4)
    return _CONVERSION_SEMAPHORE


async def _process_single_doc(
    doc_id: int,
    original_filename: str,
    stored_filename: str,
    file_size: int,
    report_dir: Path,
    pdf_dir: Path,
) -> dict:
    """Process a single document: convert to PDF, count pages, check compiled report.
    Returns a dict of fields to update on the Document row."""
    ext = Path(original_filename).suffix.lower()
    stored_path = report_dir / stored_filename
    updates: dict = {}

    if ext == ".pdf":
        # Native PDF: just count pages and check for compiled report
        updates["pdf_filename"] = stored_filename
        updates["page_count"] = await async_get_pdf_page_count(stored_path)
        updates["status"] = DocumentStatus.READY.value

        # Compiled report check — skip small files (< 5MB can't be compiled reports)
        if file_size >= 5 * 1024 * 1024:
            is_compiled = await asyncio.to_thread(is_compiled_report, stored_path)
            if is_compiled:
                updates["is_included"] = False
                updates["reasoning"] = "Auto-excluded: detected as a previously compiled report (contains TOC + multiple appendix markers)"
                logger.info(f"Auto-excluded compiled report: {original_filename}")
    else:
        # Non-PDF: convert with semaphore to limit concurrency
        async with _get_conversion_semaphore():
            pdf_result = await async_convert_to_pdf(stored_path, pdf_dir)

        if pdf_result:
            updates["pdf_filename"] = pdf_result.name
            updates["page_count"] = await async_get_pdf_page_count(pdf_result)
            updates["status"] = DocumentStatus.READY.value
        else:
            updates["status"] = DocumentStatus.ERROR.value

    return updates


async def _process_uploaded_docs(
    report_id: int,
    doc_ids: list[int],
    report_dir: Path,
    pdf_dir: Path,
):
    """Background task: convert, count pages, detect compiled reports, deduplicate."""
    try:
        db = SessionLocal()

        # Load all docs that need processing
        docs = db.query(Document).filter(Document.id.in_(doc_ids)).all()
        if not docs:
            db.close()
            return

        # Process all docs concurrently
        tasks = []
        doc_map = {}
        for doc in docs:
            doc_map[doc.id] = doc
            tasks.append(
                _process_single_doc(
                    doc_id=doc.id,
                    original_filename=doc.original_filename,
                    stored_filename=doc.stored_filename,
                    file_size=doc.file_size,
                    report_dir=report_dir,
                    pdf_dir=pdf_dir,
                )
            )

        results = await asyncio.gather(*tasks, return_exceptions=True)

        # Apply results to DB
        for doc, result in zip(docs, results):
            if isinstance(result, Exception):
                logger.error(f"Background processing failed for {doc.original_filename}: {result}")
                doc.status = DocumentStatus.ERROR.value
                continue
            for field, value in result.items():
                setattr(doc, field, value.value if hasattr(value, "value") else value)

        db.commit()

        # Deduplicate document versions
        deduplicate_documents(report_id, db)

        db.close()
        logger.info(f"Background processing complete for report {report_id}: {len(doc_ids)} docs")

    except Exception as e:
        logger.error(f"Background processing failed for report {report_id}: {e}", exc_info=True)
        try:
            db.close()
        except Exception:
            pass


@app.get("/api/reports/{report_id}/processing-status")
def get_processing_status(report_id: int, db: Session = Depends(get_db)):
    """Check how many documents are still being processed in the background."""
    total = db.query(Document).filter(Document.report_id == report_id).count()
    uploaded = db.query(Document).filter(
        Document.report_id == report_id,
        Document.status == DocumentStatus.UPLOADED.value,
    ).count()
    ready = db.query(Document).filter(
        Document.report_id == report_id,
        Document.status.in_([DocumentStatus.READY.value, DocumentStatus.CLASSIFIED.value]),
    ).count()
    errors = db.query(Document).filter(
        Document.report_id == report_id,
        Document.status == DocumentStatus.ERROR.value,
    ).count()
    return {
        "total": total,
        "pending": uploaded,
        "ready": ready,
        "errors": errors,
        "complete": uploaded == 0,
    }


@app.delete("/api/reports/{report_id}/duplicates")
async def remove_duplicate_documents(report_id: int, db: Session = Depends(get_db)):
    """Remove duplicate documents, keeping the earliest upload of each filename."""
    from sqlalchemy import func

    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "Report not found")

    subq = db.query(
        Document.original_filename,
        func.min(Document.id).label("keep_id"),
    ).filter(
        Document.report_id == report_id,
    ).group_by(Document.original_filename).subquery()

    dupes = db.query(Document).filter(
        Document.report_id == report_id,
        ~Document.id.in_(db.query(subq.c.keep_id)),
    ).all()

    count = len(dupes)
    for d in dupes:
        db.delete(d)
    db.commit()

    return {"removed": count, "remaining": db.query(Document).filter(Document.report_id == report_id).count()}


@app.post("/api/reports/{report_id}/upload-folder")
async def upload_folder(
    report_id: int,
    folder_path: str = Form(...),
    db: Session = Depends(get_db),
):
    """Upload all files from a local folder path (recursive)."""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "Report not found")

    source_dir = Path(folder_path)
    if not source_dir.exists() or not source_dir.is_dir():
        raise HTTPException(400, f"Folder not found: {folder_path}")

    report_dir = settings.UPLOAD_DIR / str(report_id) / "originals"
    report_dir.mkdir(parents=True, exist_ok=True)
    pdf_dir = settings.UPLOAD_DIR / str(report_id) / "pdfs"
    pdf_dir.mkdir(parents=True, exist_ok=True)

    uploaded = []
    skipped = []

    # Walk the folder recursively
    for file_path in sorted(source_dir.rglob("*")):
        if not file_path.is_file():
            continue

        ext = file_path.suffix.lower()

        # Skip unsupported files
        if ext in settings.SKIP_EXTENSIONS or ext not in settings.SUPPORTED_EXTENSIONS:
            skipped.append(str(file_path.name))
            continue

        # Get relative path from source folder for classification context
        try:
            rel_path = str(file_path.relative_to(source_dir))
        except ValueError:
            rel_path = file_path.name

        # Generate unique stored filename
        stored_name = f"{uuid.uuid4().hex}{ext}"
        stored_path = report_dir / stored_name

        # Copy file
        shutil.copy2(file_path, stored_path)
        file_size = stored_path.stat().st_size

        # Quick filename-based classification
        quick_class = classify_by_filename_legacy(file_path.name, rel_path)
        category = quick_class.category.value if quick_class else SectionCategory.UNCLASSIFIED.value
        subcategory = quick_class.subcategory if quick_class else None
        confidence = quick_class.confidence if quick_class else None

        doc = Document(
            report_id=report_id,
            original_filename=file_path.name,
            original_path=rel_path,
            stored_filename=stored_name,
            file_size=file_size,
            category=category,
            subcategory=subcategory,
            confidence=confidence,
            reasoning=quick_class.reasoning if quick_class else None,
            status=DocumentStatus.UPLOADED.value,
        )
        db.add(doc)
        db.flush()

        # Convert to PDF if needed
        if ext == ".pdf":
            # Copy PDF to pdfs dir too
            pdf_dest = pdf_dir / stored_name
            shutil.copy2(stored_path, pdf_dest)
            doc.pdf_filename = stored_name
            doc.page_count = get_pdf_page_count(stored_path)
            doc.status = DocumentStatus.READY.value
        else:
            pdf_result = convert_to_pdf(stored_path, pdf_dir)
            if pdf_result:
                doc.pdf_filename = pdf_result.name
                doc.page_count = get_pdf_page_count(pdf_result)
                doc.status = DocumentStatus.READY.value
            else:
                doc.status = DocumentStatus.ERROR.value

        # Bug 4: Check for compiled reports — only original PDFs
        if ext == ".pdf" and doc.status == DocumentStatus.READY.value and doc.pdf_filename:
            pdf_check_path = pdf_dir / doc.pdf_filename
            if not pdf_check_path.exists():
                pdf_check_path = stored_path
            if pdf_check_path.exists() and is_compiled_report(pdf_check_path):
                doc.is_included = False
                doc.reasoning = "Auto-excluded: detected as a previously compiled report (contains TOC + multiple appendix markers)"
                logger.info(f"Auto-excluded compiled report: {file_path.name}")

        uploaded.append(doc)

    db.commit()
    for doc in uploaded:
        db.refresh(doc)

    # Bug 5: Deduplicate document versions
    deduplicate_documents(report_id, db)

    report.status = ReportStatus.IN_PROGRESS.value
    report.updated_at = datetime.utcnow()
    db.commit()

    return {
        "uploaded": len(uploaded),
        "skipped": len(skipped),
        "skipped_files": skipped[:20],  # Show first 20
        "documents": [_doc_to_response(d) for d in uploaded],
    }


# ---- Streaming upload-folder with SSE progress ----

@app.post("/api/reports/{report_id}/upload-folder-stream")
async def upload_folder_stream(
    report_id: int,
    folder_path: str = Form(...),
    db: Session = Depends(get_db),
):
    """Upload files from a local folder with real-time SSE progress events."""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "Report not found")

    source_dir = Path(folder_path)
    if not source_dir.exists() or not source_dir.is_dir():
        raise HTTPException(400, f"Folder not found: {folder_path}")

    async def event_generator():
        # Track pipeline timing
        report.pipeline_start_time = datetime.utcnow()
        report.pipeline_end_time = None
        db.commit()

        report_dir = settings.UPLOAD_DIR / str(report_id) / "originals"
        report_dir.mkdir(parents=True, exist_ok=True)
        pdf_dir = settings.UPLOAD_DIR / str(report_id) / "pdfs"
        pdf_dir.mkdir(parents=True, exist_ok=True)

        # Collect eligible files first so we know the total
        all_files = []
        for fp in sorted(source_dir.rglob("*")):
            if not fp.is_file():
                continue
            ext = fp.suffix.lower()
            if ext in settings.SKIP_EXTENSIONS or ext not in settings.SUPPORTED_EXTENSIONS:
                continue
            all_files.append(fp)

        total = len(all_files)
        yield {"event": "progress", "data": json.dumps({"phase": "scanning", "total": total})}
        await asyncio.sleep(0)

        # ── Phase 1: Upload + Convert (all docs start as UNCLASSIFIED) ──
        uploaded = []
        errors = 0

        for i, file_path in enumerate(all_files):
            ext = file_path.suffix.lower()
            try:
                rel_path = str(file_path.relative_to(source_dir))
            except ValueError:
                rel_path = file_path.name

            yield {"event": "progress", "data": json.dumps({
                "phase": "processing",
                "current": i + 1,
                "total": total,
                "filename": file_path.name,
            })}
            await asyncio.sleep(0)

            stored_name = f"{uuid.uuid4().hex}{ext}"
            stored_path = report_dir / stored_name
            shutil.copy2(file_path, stored_path)
            file_size = stored_path.stat().st_size

            doc = Document(
                report_id=report_id,
                original_filename=file_path.name,
                original_path=rel_path,
                stored_filename=stored_name,
                file_size=file_size,
                category=SectionCategory.UNCLASSIFIED.value,
                subcategory=None,
                confidence=None,
                reasoning=None,
                status=DocumentStatus.UPLOADED.value,
            )
            db.add(doc)
            db.flush()

            # Convert to PDF
            if ext == ".pdf":
                pdf_dest = pdf_dir / stored_name
                shutil.copy2(stored_path, pdf_dest)
                doc.pdf_filename = stored_name
                doc.page_count = get_pdf_page_count(stored_path)
                doc.status = DocumentStatus.READY.value
            else:
                pdf_result = convert_to_pdf(stored_path, pdf_dir)
                if pdf_result:
                    doc.pdf_filename = pdf_result.name
                    doc.page_count = get_pdf_page_count(pdf_result)
                    doc.status = DocumentStatus.READY.value
                else:
                    doc.status = DocumentStatus.ERROR.value
                    errors += 1

            # Compiled report check — only original PDFs
            is_excluded = False
            if ext == ".pdf" and doc.status == DocumentStatus.READY.value and doc.pdf_filename:
                pdf_check_path = pdf_dir / doc.pdf_filename
                if not pdf_check_path.exists():
                    pdf_check_path = stored_path
                if pdf_check_path.exists() and is_compiled_report(pdf_check_path):
                    doc.is_included = False
                    doc.reasoning = "Auto-excluded: detected as a previously compiled report (contains TOC + multiple appendix markers)"
                    is_excluded = True

            uploaded.append(doc)

            yield {"event": "progress", "data": json.dumps({
                "phase": "processed",
                "current": i + 1,
                "total": total,
                "filename": file_path.name,
                "status": doc.status,
                "excluded": is_excluded,
                "reasoning": doc.reasoning,
            })}
            await asyncio.sleep(0)

        db.commit()
        for doc in uploaded:
            db.refresh(doc)

        # ── Phase 2: AI Classification (concurrent, queue-based) ──
        # Gather docs that need classification (non-error, included)
        to_classify = []
        for doc in uploaded:
            if doc.status == DocumentStatus.ERROR.value:
                continue
            if not doc.is_included:
                continue
            if doc.pdf_filename:
                doc_pdf_path = pdf_dir / doc.pdf_filename
                if not doc_pdf_path.exists():
                    doc_pdf_path = report_dir / doc.stored_filename
            else:
                doc_pdf_path = report_dir / doc.stored_filename
            to_classify.append((doc.id, doc_pdf_path, doc.original_filename, doc.original_path or ""))

        classify_total = len(to_classify)
        if classify_total > 0:
            yield {"event": "progress", "data": json.dumps({
                "phase": "classifying",
                "current": 0,
                "total": classify_total,
            })}
            await asyncio.sleep(0)

            result_queue: asyncio.Queue = asyncio.Queue()
            classify_task = asyncio.create_task(
                classify_all_documents_queued(to_classify, result_queue)
            )

            # Build doc lookup for fast updates
            doc_map = {doc.id: doc for doc in uploaded}

            while True:
                item = await result_queue.get()
                if item is None:
                    break  # sentinel — all done

                doc_id = item["doc_id"]
                result = item["result"]
                fname = item["filename"]
                current = item["current"]
                ctotal = item["total"]

                doc = doc_map.get(doc_id)
                if doc:
                    # Apply preference rules (rev/marked, property profile sort)
                    all_fnames = [d.original_filename for d in doc_map.values()]
                    from classifier import apply_preference_rules
                    result, excludes = apply_preference_rules(result, fname, all_fnames)
                    for exc_name in excludes:
                        for d in doc_map.values():
                            if d.original_filename.lower() == exc_name.lower() and d.is_included:
                                d.is_included = False
                                d.reasoning = (d.reasoning or "") + " [Superseded by revised version]"

                    doc.category = result.category.value
                    doc.subcategory = result.subcategory
                    doc.confidence = result.confidence
                    doc.reasoning = result.reasoning
                    doc.status = DocumentStatus.CLASSIFIED.value
                    if result.sort_order is not None:
                        doc.sort_order = result.sort_order

                yield {"event": "progress", "data": json.dumps({
                    "phase": "classified",
                    "current": current,
                    "total": ctotal,
                    "filename": fname,
                    "category": result.category.value,
                    "confidence": result.confidence,
                    "reasoning": result.reasoning,
                })}
                await asyncio.sleep(0)

            # Wait for task to fully complete (handle any exceptions)
            await classify_task

            db.commit()

        # Deduplication
        deduplicate_documents(report_id, db)

        # Assembly validation
        yield {"event": "progress", "data": json.dumps({
            "phase": "validating",
            "current": 0,
            "total": 0,
        })}
        await asyncio.sleep(0)

        validation = await validate_assembly(report_id, db)

        # Report Director — review manifest and flag curation recommendations
        yield {"event": "progress", "data": json.dumps({
            "phase": "directing",
            "current": 0,
            "total": 0,
        })}
        await asyncio.sleep(0)

        director_result = await run_report_director(report_id, db)

        report.status = ReportStatus.IN_PROGRESS.value
        report.updated_at = datetime.utcnow()
        db.commit()

        yield {"event": "complete", "data": json.dumps({
            "uploaded": len(uploaded),
            "classified": classify_total,
            "errors": errors,
            "total": total,
            "validation_applied": validation.get("applied", 0),
            "director_health": director_result.get("health", "unknown"),
            "director_exclude_count": director_result.get("exclude_count", 0),
            "director_estimated_pages": director_result.get("estimated_pages"),
        })}

    return EventSourceResponse(event_generator())


# ---- Streaming classification with SSE progress ----

@app.post("/api/reports/{report_id}/classify-stream")
async def classify_documents_stream(report_id: int, db: Session = Depends(get_db)):
    """Run AI classification with real-time SSE progress events."""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "Report not found")

    async def event_generator():
        docs = db.query(Document).filter(
            Document.report_id == report_id,
        ).all()

        to_classify_docs = [
            d for d in docs
            if d.category == SectionCategory.UNCLASSIFIED.value
            or (d.confidence is not None and d.confidence < 0.85)
        ]

        total = len(to_classify_docs)
        yield {"event": "progress", "data": json.dumps({"phase": "starting", "total": total})}
        await asyncio.sleep(0)

        if total > 0:
            pdf_dir = settings.UPLOAD_DIR / str(report_id) / "pdfs"
            originals_dir = settings.UPLOAD_DIR / str(report_id) / "originals"

            # Build classification input list
            classify_input = []
            for doc in to_classify_docs:
                doc.status = DocumentStatus.CLASSIFYING.value
                if doc.pdf_filename:
                    file_path = pdf_dir / doc.pdf_filename
                else:
                    file_path = originals_dir / doc.stored_filename
                classify_input.append((doc.id, file_path, doc.original_filename, doc.original_path or ""))
            db.commit()

            doc_map = {doc.id: doc for doc in to_classify_docs}

            result_queue: asyncio.Queue = asyncio.Queue()
            classify_task = asyncio.create_task(
                classify_all_documents_queued(classify_input, result_queue)
            )

            classified_count = 0
            while True:
                item = await result_queue.get()
                if item is None:
                    break

                doc_id = item["doc_id"]
                result = item["result"]
                fname = item["filename"]
                current = item["current"]
                ctotal = item["total"]

                doc = doc_map.get(doc_id)
                if doc:
                    # Apply preference rules
                    all_fnames = [d.original_filename for d in doc_map.values()]
                    from classifier import apply_preference_rules
                    result, excludes = apply_preference_rules(result, fname, all_fnames)
                    for exc_name in excludes:
                        for d in doc_map.values():
                            if d.original_filename.lower() == exc_name.lower() and d.is_included:
                                d.is_included = False
                                d.reasoning = (d.reasoning or "") + " [Superseded by revised version]"

                    doc.category = result.category.value
                    doc.subcategory = result.subcategory
                    doc.confidence = result.confidence
                    doc.reasoning = result.reasoning
                    doc.status = DocumentStatus.CLASSIFIED.value
                    if result.sort_order is not None:
                        doc.sort_order = result.sort_order
                    classified_count += 1

                yield {"event": "progress", "data": json.dumps({
                    "phase": "classified",
                    "current": current,
                    "total": ctotal,
                    "filename": fname,
                    "category": result.category.value,
                    "confidence": result.confidence,
                    "reasoning": result.reasoning,
                })}
                await asyncio.sleep(0)

            await classify_task
            db.commit()

        # Run assembly validation after classification
        yield {"event": "progress", "data": json.dumps({
            "phase": "validating",
            "current": 0,
            "total": 0,
        })}
        await asyncio.sleep(0)

        validation = await validate_assembly(report_id, db)

        yield {"event": "complete", "data": json.dumps({
            "classified": total,
            "validation_applied": validation.get("applied", 0),
            "validation_skipped": validation.get("skipped", False),
        })}

    return EventSourceResponse(event_generator())


# ---- Assembly Validation ----

@app.post("/api/reports/{report_id}/validate-assembly")
async def validate_assembly_endpoint(report_id: int, db: Session = Depends(get_db)):
    """Run AI validation on the document manifest to catch cross-document misclassifications."""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "Report not found")

    result = await validate_assembly(report_id, db)
    return {
        "status": "ok",
        "applied": result.get("applied", 0),
        "flagged": result.get("flagged", 0),
        "skipped": result.get("skipped", False),
    }


# ---- Report Director ----

@app.post("/api/reports/{report_id}/director")
async def run_director_endpoint(report_id: int, db: Session = Depends(get_db)):
    """Run the Report Director to review manifest and recommend curation."""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "Report not found")
    result = await run_report_director(report_id, db)
    return result


@app.post("/api/reports/{report_id}/apply-director")
async def apply_director_recommendations(report_id: int, db: Session = Depends(get_db)):
    """Apply all director exclude recommendations — bulk set is_included=false."""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "Report not found")

    docs = db.query(Document).filter(
        Document.report_id == report_id,
        Document.is_included == True,
    ).all()

    excluded = 0
    for doc in docs:
        if doc.reasoning and "[DIRECTOR: recommend exclude" in (doc.reasoning or ""):
            doc.is_included = False
            excluded += 1

    if excluded > 0:
        db.commit()

    remaining_docs = db.query(Document).filter(
        Document.report_id == report_id,
        Document.is_included == True,
    ).count()
    remaining_pages = sum(
        d.page_count or 0 for d in db.query(Document).filter(
            Document.report_id == report_id,
            Document.is_included == True,
        ).all()
    )

    return {
        "excluded": excluded,
        "remaining_docs": remaining_docs,
        "remaining_pages": remaining_pages,
    }


# ---- Bug 6: Reprocess errors ----

@app.post("/api/reports/{report_id}/reprocess-errors")
async def reprocess_errors(report_id: int, db: Session = Depends(get_db)):
    """Re-attempt PDF conversion on all documents with error status."""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "Report not found")

    error_docs = db.query(Document).filter(
        Document.report_id == report_id,
        Document.status == DocumentStatus.ERROR.value,
    ).all()

    if not error_docs:
        return {"fixed": 0, "remaining_errors": 0}

    pdf_dir = settings.UPLOAD_DIR / str(report_id) / "pdfs"
    pdf_dir.mkdir(parents=True, exist_ok=True)
    originals_dir = settings.UPLOAD_DIR / str(report_id) / "originals"

    fixed = 0
    for doc in error_docs:
        stored_path = originals_dir / doc.stored_filename
        if not stored_path.exists():
            continue

        pdf_result = convert_to_pdf(stored_path, pdf_dir)
        if pdf_result:
            doc.pdf_filename = pdf_result.name
            doc.page_count = get_pdf_page_count(pdf_result)
            doc.status = DocumentStatus.READY.value
            fixed += 1

    db.commit()

    remaining = db.query(Document).filter(
        Document.report_id == report_id,
        Document.status == DocumentStatus.ERROR.value,
    ).count()

    return {"fixed": fixed, "remaining_errors": remaining}


# ---- AI Classification ----

@app.post("/api/reports/{report_id}/classify")
async def classify_all_documents(report_id: int, db: Session = Depends(get_db)):
    """Run AI classification on all unclassified or low-confidence documents."""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "Report not found")

    docs = db.query(Document).filter(
        Document.report_id == report_id,
    ).all()

    # Only re-classify docs that are unclassified or low confidence
    to_classify = [
        d for d in docs
        if d.category == SectionCategory.UNCLASSIFIED.value
        or (d.confidence is not None and d.confidence < 0.85)
    ]

    results = []
    for doc in to_classify:
        doc.status = DocumentStatus.CLASSIFYING.value
        db.commit()

        # Determine file path
        if doc.pdf_filename:
            file_path = settings.UPLOAD_DIR / str(report_id) / "pdfs" / doc.pdf_filename
        else:
            file_path = settings.UPLOAD_DIR / str(report_id) / "originals" / doc.stored_filename

        result = await classify_document(
            file_path,
            doc.original_filename,
            doc.original_path or "",
        )

        # Apply preference rules
        all_fnames = [d.original_filename for d in docs]
        from classifier import apply_preference_rules
        result, excludes = apply_preference_rules(result, doc.original_filename, all_fnames)
        for exc_name in excludes:
            for d in docs:
                if d.original_filename.lower() == exc_name.lower() and d.is_included:
                    d.is_included = False
                    d.reasoning = (d.reasoning or "") + " [Superseded by revised version]"

        doc.category = result.category.value
        doc.subcategory = result.subcategory
        doc.confidence = result.confidence
        doc.reasoning = result.reasoning
        doc.status = DocumentStatus.CLASSIFIED.value
        if result.sort_order is not None:
            doc.sort_order = result.sort_order
        db.commit()

        results.append({
            "document_id": doc.id,
            "filename": doc.original_filename,
            "category": result.category.value,
            "subcategory": result.subcategory,
            "confidence": result.confidence,
            "reasoning": result.reasoning,
        })

    return {"classified": len(results), "results": results}


@app.post("/api/reports/{report_id}/auto-name")
async def auto_name_report(report_id: int, db: Session = Depends(get_db)):
    """Extract project number and address from the Cover/Write-Up document using AI."""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "Report not found")

    # Find the Cover/Write-Up document
    cover_docs = db.query(Document).filter(
        Document.report_id == report_id,
        Document.category == SectionCategory.COVER_WRITEUP.value,
        Document.is_included == True,
    ).all()

    if not cover_docs:
        return {"status": "skipped", "reason": "No Cover/Write-Up document found"}

    # Extract text from the first cover doc
    doc = cover_docs[0]
    if doc.pdf_filename:
        file_path = settings.UPLOAD_DIR / str(report_id) / "pdfs" / doc.pdf_filename
        if not file_path.exists():
            file_path = settings.UPLOAD_DIR / str(report_id) / "originals" / doc.pdf_filename
    else:
        file_path = settings.UPLOAD_DIR / str(report_id) / "originals" / doc.stored_filename

    if not file_path.exists():
        return {"status": "skipped", "reason": "Cover document file not found"}

    # Import text extraction from classifier
    from classifier import _extract_text_from_pdf, _extract_text_from_docx

    ext = file_path.suffix.lower()
    if ext == ".pdf":
        text = _extract_text_from_pdf(file_path, max_pages=5, max_chars=5000)
    elif ext in (".docx", ".doc"):
        text = _extract_text_from_docx(file_path, max_chars=5000)
    else:
        return {"status": "skipped", "reason": "Cover document is not a PDF or DOCX"}

    if len(text.strip()) < 20:
        return {"status": "skipped", "reason": "Could not extract enough text from cover document"}

    # Ask Ollama to extract project info
    extraction_prompt = """Extract the following from this Phase I Environmental Site Assessment report. Use both the filename and text content to find the information. Return ONLY a JSON object with these fields:

{
  "project_number": "the project number/reference number (e.g. '6384674' or 'P-12345')",
  "address": "the property/site address (street address, city, state, zip)",
  "name": "a short descriptive name for the report (e.g. 'Phase I ESA - 123 Main St, Springfield')"
}

If a field cannot be determined, use null. Do NOT guess — only extract what is clearly stated."""

    try:
        prompt_text = f"{extraction_prompt}\n\nFilename: {doc.original_filename}\n\nReport text:\n{text[:4000]}"
        async with httpx.AsyncClient(timeout=60.0) as client:
            resp = await client.post(
                f"{settings.OLLAMA_URL}/api/generate",
                json={
                    "model": settings.OLLAMA_MODEL,
                    "prompt": prompt_text,
                    "stream": False,
                    "format": "json",
                },
            )
            resp.raise_for_status()
            raw = resp.json()["response"].strip()

            # Parse JSON (strip markdown fences if present)
            import re as _re
            if raw.startswith("```"):
                raw = _re.sub(r"```(?:json)?\s*", "", raw)
                raw = raw.rstrip("`").strip()

            extracted = json.loads(raw)

        updated_fields = {}

        if extracted.get("project_number") and not report.project_number:
            report.project_number = str(extracted["project_number"])
            updated_fields["project_number"] = report.project_number

        if extracted.get("address") and not report.address:
            report.address = str(extracted["address"])
            updated_fields["address"] = report.address

        if extracted.get("name") and report.name.startswith("New Report"):
            report.name = str(extracted["name"])
            updated_fields["name"] = report.name

        if updated_fields:
            report.updated_at = datetime.utcnow()
            db.commit()

        return {"status": "ok", "updated": updated_fields, "extracted": extracted}

    except Exception as e:
        logger.error(f"Auto-name failed for report {report_id}: {e}")
        return {"status": "error", "reason": str(e)}


# ---- Document Management ----

# Batch endpoint MUST come before /{doc_id} routes to avoid "batch" matching as an int
@app.put("/api/reports/{report_id}/documents/batch")
def batch_update_documents(report_id: int, request: BatchUpdateRequest, db: Session = Depends(get_db)):
    """Batch update documents — move to section and/or change inclusion."""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "Report not found")

    docs = db.query(Document).filter(
        Document.report_id == report_id,
        Document.id.in_(request.document_ids),
    ).all()

    for doc in docs:
        if request.category is not None:
            doc.category = request.category.value
            doc.confidence = 1.0
            doc.reasoning = "Batch update by user"
        if request.is_included is not None:
            doc.is_included = request.is_included

    db.commit()
    return {"updated": len(docs)}


@app.get("/api/reports/{report_id}/documents", response_model=list[DocumentResponse])
def list_documents(report_id: int, db: Session = Depends(get_db)):
    docs = db.query(Document).filter(
        Document.report_id == report_id
    ).order_by(Document.category, Document.sort_order, Document.id).all()
    return [_doc_to_response(d) for d in docs]


@app.put("/api/reports/{report_id}/documents/{doc_id}", response_model=DocumentResponse)
def update_document(
    report_id: int,
    doc_id: int,
    data: DocumentUpdate,
    db: Session = Depends(get_db),
):
    doc = db.query(Document).filter(
        Document.id == doc_id,
        Document.report_id == report_id,
    ).first()
    if not doc:
        raise HTTPException(404, "Document not found")

    for field, value in data.dict(exclude_unset=True).items():
        if value is not None:
            setattr(doc, field, value.value if hasattr(value, "value") else value)

    # If manually reclassified, set confidence to 1.0
    if data.category is not None:
        doc.confidence = 1.0
        doc.reasoning = "Manually classified by user"

    db.commit()
    db.refresh(doc)
    return _doc_to_response(doc)


@app.delete("/api/reports/{report_id}/documents/{doc_id}")
def delete_document(report_id: int, doc_id: int, db: Session = Depends(get_db)):
    doc = db.query(Document).filter(
        Document.id == doc_id,
        Document.report_id == report_id,
    ).first()
    if not doc:
        raise HTTPException(404, "Document not found")

    # Delete files
    originals_dir = settings.UPLOAD_DIR / str(report_id) / "originals"
    pdfs_dir = settings.UPLOAD_DIR / str(report_id) / "pdfs"

    (originals_dir / doc.stored_filename).unlink(missing_ok=True)
    if doc.pdf_filename:
        (pdfs_dir / doc.pdf_filename).unlink(missing_ok=True)

    db.delete(doc)
    db.commit()
    return {"status": "deleted"}


@app.put("/api/reports/{report_id}/reorder")
def reorder_documents(
    report_id: int,
    data: ReorderRequest,
    db: Session = Depends(get_db),
):
    """Reorder documents within a section."""
    for i, doc_id in enumerate(data.document_ids):
        doc = db.query(Document).filter(
            Document.id == doc_id,
            Document.report_id == report_id,
        ).first()
        if doc:
            doc.sort_order = i
            doc.category = data.category.value

    db.commit()
    return {"status": "reordered"}


# ---- Document Preview ----

@app.get("/api/reports/{report_id}/documents/{doc_id}/preview")
def preview_document(report_id: int, doc_id: int, db: Session = Depends(get_db)):
    doc = db.query(Document).filter(
        Document.id == doc_id,
        Document.report_id == report_id,
    ).first()
    if not doc:
        raise HTTPException(404, "Document not found")

    if doc.pdf_filename:
        pdf_path = settings.UPLOAD_DIR / str(report_id) / "pdfs" / doc.pdf_filename
        if not pdf_path.exists():
            # Try originals dir for PDFs
            pdf_path = settings.UPLOAD_DIR / str(report_id) / "originals" / doc.pdf_filename
    else:
        pdf_path = settings.UPLOAD_DIR / str(report_id) / "originals" / doc.stored_filename

    if not pdf_path.exists():
        raise HTTPException(404, "PDF file not found")

    return FileResponse(
        pdf_path,
        media_type="application/pdf",
        content_disposition_type="inline",
    )


# ---- Report Assembly ----

@app.get("/api/reports/{report_id}/preflight")
async def preflight_check(report_id: int, db: Session = Depends(get_db)):
    """Pre-flight check before assembly — returns warnings and errors."""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "Report not found")

    docs = db.query(Document).filter(
        Document.report_id == report_id,
        Document.is_included == True,
        Document.status.in_([DocumentStatus.READY.value, DocumentStatus.CLASSIFIED.value]),
    ).all()

    warnings = []
    errors = []

    # Group by section
    by_section: dict[str, list] = {}
    for d in docs:
        by_section.setdefault(d.category, []).append(d)

    # Must have Cover/Write-Up
    if SectionCategory.COVER_WRITEUP.value not in by_section:
        errors.append("Missing Cover / Write-Up — report has no main body document")

    # Check for unclassified
    unclassified = by_section.get(SectionCategory.UNCLASSIFIED.value, [])
    if unclassified:
        warnings.append(f"{len(unclassified)} document{'s' if len(unclassified) != 1 else ''} still unclassified")

    # Check empty sections that should typically have content
    expected_sections = [
        (SectionCategory.APPENDIX_A.value, "Appendix A — Maps & Plot Plan"),
        (SectionCategory.APPENDIX_B.value, "Appendix B — Site Photographs"),
        (SectionCategory.APPENDIX_C.value, "Appendix C — Database Report"),
        (SectionCategory.APPENDIX_D.value, "Appendix D — Historical Records"),
    ]
    empty_sections = [label for val, label in expected_sections if val not in by_section]
    if empty_sections:
        warnings.append(f"Empty sections: {', '.join(empty_sections)}")

    # Page count checks
    total_pages = sum(d.page_count or 0 for d in docs)
    if total_pages > 2000:
        errors.append(f"Report has {total_pages} pages — likely contains duplicates. Check documents before assembling.")
    elif total_pages > 500:
        warnings.append(f"Report has {total_pages} pages — verify this is expected")

    sections_filled = len([s for s in settings.SECTION_ORDER if s in by_section])
    sections_empty = len(settings.SECTION_ORDER) - sections_filled

    return {
        "can_assemble": len(errors) == 0 and len(docs) > 0,
        "warnings": warnings,
        "errors": errors,
        "stats": {
            "total_pages": total_pages,
            "total_docs": len(docs),
            "sections_filled": sections_filled,
            "sections_empty": sections_empty,
        },
    }


@app.post("/api/reports/{report_id}/assemble")
async def assemble_report_endpoint(
    report_id: int,
    data: AssembleRequest = None,
    db: Session = Depends(get_db),
):
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "Report not found")

    docs = db.query(Document).filter(
        Document.report_id == report_id,
        Document.is_included == True,
        Document.status.in_([DocumentStatus.READY.value, DocumentStatus.CLASSIFIED.value]),
    ).order_by(Document.category, Document.sort_order, Document.id).all()

    if not docs:
        raise HTTPException(400, "No documents ready for assembly")

    # Build document list for assembler
    pdf_dir = settings.UPLOAD_DIR / str(report_id) / "pdfs"
    originals_dir = settings.UPLOAD_DIR / str(report_id) / "originals"
    output_dir = settings.UPLOAD_DIR / str(report_id) / "output"
    output_dir.mkdir(parents=True, exist_ok=True)

    doc_list = []
    skipped = []
    for doc in docs:
        if doc.pdf_filename:
            pdf_path = pdf_dir / doc.pdf_filename
            if not pdf_path.exists():
                pdf_path = originals_dir / doc.pdf_filename
        else:
            pdf_path = originals_dir / doc.stored_filename

        if pdf_path.exists() and pdf_path.suffix.lower() == ".pdf":
            doc_list.append({
                "doc_id": doc.id,
                "pdf_path": pdf_path,
                "category": doc.category,
                "subcategory": doc.subcategory,
                "sort_order": doc.sort_order,
                "original_filename": doc.original_filename,
                "original_path": doc.original_path or "",
            })
        else:
            skipped.append(doc.original_filename)
            logger.warning(f"Assembly skipped {doc.original_filename}: PDF not found at {pdf_path}")

    if not doc_list:
        raise HTTPException(400, f"No valid PDF files found for assembly ({len(skipped)} docs skipped — missing or unconverted)")

    # Assemble
    output_filename = f"{report.project_number or report.name}_assembled.pdf"
    # Sanitize filename
    output_filename = re.sub(r'[^\w\s\-.]', '_', output_filename)
    output_path = output_dir / output_filename

    try:
        result = assemble_report(
            documents=doc_list,
            output_path=output_path,
            has_reliance_letter=report.has_reliance_letter,
        )
    except Exception as e:
        logger.error(f"Assembly failed for report {report_id}: {e}", exc_info=True)
        raise HTTPException(500, f"Assembly failed: {str(e)}")

    # Compress if requested
    compressed_size = None
    if data and data.compression:
        compressed_path = output_dir / f"{report.project_number or report.name}_compressed.pdf"
        comp_result = compress_pdf(output_path, compressed_path, quality=data.compression)
        compressed_size = comp_result["compressed_size"]
        report.compressed_size = compressed_size

    report.assembled_filename = output_filename
    report.assembled_size = result["file_size"]
    report.status = ReportStatus.DONE.value
    report.updated_at = datetime.utcnow()
    # Save manifest and pipeline end time
    if result.get("document_manifest"):
        report.manifest_json = json.dumps(result["document_manifest"])
    if report.pipeline_start_time and not report.pipeline_end_time:
        report.pipeline_end_time = datetime.utcnow()
    db.commit()

    # Check for page count anomalies
    warnings = []
    expected_pages = sum(d.page_count or 0 for d in docs)
    if expected_pages > 0 and result["total_pages"] > expected_pages * 1.1:
        warnings.append(f"Assembled {result['total_pages']} pages but expected ~{expected_pages} — check for duplicates")

    return {
        "status": "assembled",
        "total_pages": result["total_pages"],
        "total_documents": result["total_documents"],
        "file_size": result["file_size"],
        "file_size_display": get_file_size_display(result["file_size"]),
        "compressed_size": compressed_size,
        "compressed_size_display": get_file_size_display(compressed_size) if compressed_size else None,
        "section_pages": result["section_pages"],
        "document_manifest": result.get("document_manifest", []),
        "errors": result["errors"],
        "warnings": warnings,
    }


@app.get("/api/reports/{report_id}/preview")
def preview_assembled_report(report_id: int, db: Session = Depends(get_db)):
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report or not report.assembled_filename:
        raise HTTPException(404, "Assembled report not found")

    output_dir = settings.UPLOAD_DIR / str(report_id) / "output"
    pdf_path = output_dir / report.assembled_filename
    if not pdf_path.exists():
        raise HTTPException(404, "Assembled PDF file not found")

    return FileResponse(
        pdf_path,
        media_type="application/pdf",
        content_disposition_type="inline",
    )


@app.get("/api/reports/{report_id}/assembled/page/{page_num}")
def get_assembled_page(
    report_id: int,
    page_num: int,
    width: int = Query(1600, ge=200, le=4000),
    db: Session = Depends(get_db),
):
    """Render a single page of the assembled PDF as a PNG image."""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report or not report.assembled_filename:
        raise HTTPException(404, "Assembled report not found")

    output_dir = settings.UPLOAD_DIR / str(report_id) / "output"
    pdf_path = output_dir / report.assembled_filename
    if not pdf_path.exists():
        raise HTTPException(404, "Assembled PDF file not found")

    # Cache directory for rendered pages
    cache_dir = output_dir / "page_cache" / f"w{width}"
    cache_dir.mkdir(parents=True, exist_ok=True)
    cache_file = cache_dir / f"page_{page_num}.png"

    # Check if assembled PDF is newer than cache
    if cache_file.exists() and cache_file.stat().st_mtime >= pdf_path.stat().st_mtime:
        return FileResponse(
            cache_file,
            media_type="image/png",
            headers={"Cache-Control": "public, max-age=3600"},
        )

    # Render the page using pypdfium2
    try:
        import pypdfium2 as pdfium

        pdf = pdfium.PdfDocument(str(pdf_path))
        if page_num < 1 or page_num > len(pdf):
            pdf.close()
            raise HTTPException(400, f"Page {page_num} out of range (1-{len(pdf)})")

        page = pdf[page_num - 1]  # 0-indexed internally
        # Scale to requested width
        page_width = page.get_width()
        scale = width / page_width
        bitmap = page.render(scale=scale)
        pil_image = bitmap.to_pil()
        pil_image.save(str(cache_file), "PNG", optimize=True)
        pdf.close()
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to render page {page_num} of report {report_id}: {e}")
        raise HTTPException(500, f"Failed to render page: {e}")

    return FileResponse(
        cache_file,
        media_type="image/png",
        headers={"Cache-Control": "public, max-age=3600"},
    )


@app.get("/api/reports/{report_id}/download")
def download_report(
    report_id: int,
    compressed: bool = Query(False),
    db: Session = Depends(get_db),
):
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "Report not found")

    output_dir = settings.UPLOAD_DIR / str(report_id) / "output"

    if compressed:
        # Look for compressed version
        compressed_name = f"{report.project_number or report.name}_compressed.pdf"
        pdf_path = output_dir / compressed_name
        if not pdf_path.exists() and report.assembled_filename:
            pdf_path = output_dir / report.assembled_filename
    else:
        if not report.assembled_filename:
            raise HTTPException(404, "No assembled report")
        pdf_path = output_dir / report.assembled_filename

    if not pdf_path.exists():
        raise HTTPException(404, "PDF not found")

    return FileResponse(
        pdf_path,
        media_type="application/pdf",
        filename=pdf_path.name,
    )


@app.post("/api/reports/{report_id}/compress")
def compress_report(
    report_id: int,
    data: CompressRequest,
    db: Session = Depends(get_db),
):
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report or not report.assembled_filename:
        raise HTTPException(404, "No assembled report to compress")

    output_dir = settings.UPLOAD_DIR / str(report_id) / "output"
    input_path = output_dir / report.assembled_filename
    compressed_name = f"{report.project_number or report.name}_compressed.pdf"
    output_path = output_dir / compressed_name

    result = compress_pdf(
        input_path,
        output_path,
        quality=data.quality,
        target_size_mb=data.target_size_mb,
    )

    report.compressed_size = result["compressed_size"]
    report.updated_at = datetime.utcnow()
    db.commit()

    return {
        "original_size": result["original_size"],
        "original_size_display": get_file_size_display(result["original_size"]),
        "compressed_size": result["compressed_size"],
        "compressed_size_display": get_file_size_display(result["compressed_size"]),
        "reduction_pct": result["reduction_pct"],
    }


# ---- Chat / Command Bar ----

@app.post("/api/reports/{report_id}/chat")
async def chat_with_report(report_id: int, request: ChatRequest, db: Session = Depends(get_db)):
    """Send a natural language message to manage the report."""
    from chat import process_message
    response = await process_message(report_id, request.message, db)

    # Handle deferred actions
    for result in response.results:
        if isinstance(result, dict) and result.get("deferred"):
            action = result.get("action")
            params = result.get("params", {})
            if action == "assemble":
                report = db.query(Report).filter(Report.id == report_id).first()
                if report:
                    docs = db.query(Document).filter(
                        Document.report_id == report_id,
                        Document.is_included == True,
                    ).all()
                    doc_dicts = []
                    for doc in docs:
                        pdf_path = settings.UPLOAD_DIR / str(report_id) / "pdfs" / (doc.pdf_filename or doc.stored_filename)
                        doc_dicts.append({
                            "doc_id": doc.id,
                            "pdf_path": str(pdf_path),
                            "category": doc.category,
                            "subcategory": doc.subcategory,
                            "sort_order": doc.sort_order,
                            "original_filename": doc.original_filename,
                            "original_path": doc.original_path or "",
                            "page_count": doc.page_count,
                        })
                    output_dir = settings.UPLOAD_DIR / str(report_id) / "output"
                    output_dir.mkdir(parents=True, exist_ok=True)
                    output_path = output_dir / f"report_{report_id}.pdf"
                    assembly_result = assemble_report(doc_dicts, output_path, report.has_reliance_letter)
                    report.assembled_filename = output_path.name
                    report.assembled_size = assembly_result["file_size"]
                    report.status = ReportStatus.DONE.value
                    db.commit()

    return {
        "message": response.message,
        "actions": [{"action": a.action, "params": a.params} for a in response.actions],
        "results": response.results,
        "needs_confirmation": response.needs_confirmation,
        "affected_count": response.affected_count,
    }


@app.post("/api/reports/{report_id}/undo")
def undo_last(report_id: int, db: Session = Depends(get_db)):
    """Undo the last chat action."""
    from chat import undo_last_action
    return undo_last_action(report_id, db)


@app.get("/api/reports/{report_id}/chat-history")
def get_chat_history(report_id: int, db: Session = Depends(get_db)):
    """Get chat conversation history."""
    messages = db.query(ChatMessage).filter(
        ChatMessage.report_id == report_id,
    ).order_by(ChatMessage.created_at).all()
    return [{
        "id": m.id,
        "report_id": m.report_id,
        "role": m.role,
        "content": m.content,
        "actions": json.loads(m.actions_json) if m.actions_json else None,
        "created_at": m.created_at.isoformat() if m.created_at else None,
    } for m in messages]


@app.get("/api/reports/{report_id}/suggestions")
def get_suggestions(report_id: int, db: Session = Depends(get_db)):
    """Get contextual command suggestions."""
    from chat import get_contextual_suggestions
    return {"suggestions": get_contextual_suggestions(report_id, db)}


# ---- Split for Email ----

@app.post("/api/reports/{report_id}/split")
def split_report(report_id: int, max_size_mb: float = Query(20.0), db: Session = Depends(get_db)):
    """Split assembled PDF into email-sized parts."""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report or not report.assembled_filename:
        raise HTTPException(404, "No assembled report found")

    assembled_path = settings.UPLOAD_DIR / str(report_id) / "output" / report.assembled_filename
    if not assembled_path.exists():
        raise HTTPException(404, "Assembled PDF file not found")

    from splitter import split_pdf
    parts = split_pdf(assembled_path, max_size_mb)

    return {
        "parts": [{
            "part_number": p["part_number"],
            "filename": p["filename"],
            "start_page": p["start_page"],
            "end_page": p["end_page"],
            "page_count": p["page_count"],
            "file_size": p["file_size"],
        } for p in parts],
        "total_parts": len(parts),
    }


@app.get("/api/reports/{report_id}/split/{part_num}")
def download_split_part(report_id: int, part_num: int, db: Session = Depends(get_db)):
    """Download an individual split part."""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report or not report.assembled_filename:
        raise HTTPException(404, "No assembled report found")

    stem = Path(report.assembled_filename).stem
    suffix = Path(report.assembled_filename).suffix
    part_filename = f"{stem}_part{part_num}{suffix}"
    part_path = settings.UPLOAD_DIR / str(report_id) / "output" / "split" / part_filename

    if not part_path.exists():
        raise HTTPException(404, f"Split part {part_num} not found")

    return FileResponse(
        path=str(part_path),
        filename=part_filename,
        media_type="application/pdf",
    )


# ---- Text Editing ----

@app.post("/api/reports/{report_id}/documents/{doc_id}/text-replace")
async def text_replace(report_id: int, doc_id: int, request: TextReplaceRequest, db: Session = Depends(get_db)):
    """Find/replace text in a DOCX document, then re-convert to PDF."""
    doc = db.query(Document).filter(Document.id == doc_id, Document.report_id == report_id).first()
    if not doc:
        raise HTTPException(404, "Document not found")

    original_path = settings.UPLOAD_DIR / str(report_id) / "originals" / doc.stored_filename
    if not original_path.exists() or not doc.stored_filename.lower().endswith((".docx",)):
        raise HTTPException(400, "Text replace only works on DOCX files")

    try:
        from docx import Document as DocxDocument
        docx_doc = DocxDocument(str(original_path))
        count = 0

        for paragraph in docx_doc.paragraphs:
            if request.find in paragraph.text:
                for run in paragraph.runs:
                    if request.find in run.text:
                        run.text = run.text.replace(request.find, request.replace)
                        count += 1

        for table in docx_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if request.find in paragraph.text:
                            for run in paragraph.runs:
                                if request.find in run.text:
                                    run.text = run.text.replace(request.find, request.replace)
                                    count += 1

        docx_doc.save(str(original_path))

        # Re-convert to PDF
        pdf_dir = settings.UPLOAD_DIR / str(report_id) / "pdfs"
        pdf_dir.mkdir(parents=True, exist_ok=True)
        pdf_path = await async_convert_to_pdf(original_path, pdf_dir)
        if pdf_path:
            doc.pdf_filename = pdf_path.name
            page_count = await async_get_pdf_page_count(pdf_path)
            if page_count is not None:
                doc.page_count = page_count
            db.commit()

        return {"status": "ok", "replacements": count}
    except Exception as e:
        raise HTTPException(500, f"Text replace failed: {str(e)}")


@app.post("/api/reports/{report_id}/documents/{doc_id}/delete-pages")
async def delete_pages(report_id: int, doc_id: int, request: DeletePagesRequest, db: Session = Depends(get_db)):
    """Remove specific pages from a PDF document."""
    doc = db.query(Document).filter(Document.id == doc_id, Document.report_id == report_id).first()
    if not doc:
        raise HTTPException(404, "Document not found")

    pdf_path = settings.UPLOAD_DIR / str(report_id) / "pdfs" / (doc.pdf_filename or doc.stored_filename)
    if not pdf_path.exists():
        raise HTTPException(404, "PDF file not found")

    try:
        from pypdf import PdfReader, PdfWriter
        reader = PdfReader(str(pdf_path))
        writer = PdfWriter()

        pages_to_remove = set(request.pages)
        for i, page in enumerate(reader.pages):
            if i not in pages_to_remove:
                writer.add_page(page)

        with open(pdf_path, "wb") as f:
            writer.write(f)

        doc.page_count = len(reader.pages) - len(pages_to_remove)
        db.commit()

        return {"status": "ok", "remaining_pages": doc.page_count}
    except Exception as e:
        raise HTTPException(500, f"Delete pages failed: {str(e)}")


# ---- DOCX Inline Editing ----

@app.get("/api/reports/{report_id}/documents/{doc_id}/docx-content")
def get_docx_content(report_id: int, doc_id: int, db: Session = Depends(get_db)):
    """Extract structured content from a DOCX document for inline editing."""
    doc = db.query(Document).filter(Document.id == doc_id, Document.report_id == report_id).first()
    if not doc:
        raise HTTPException(404, "Document not found")

    if not doc.stored_filename.lower().endswith((".docx",)):
        return DocxContentResponse(is_docx=False)

    original_path = settings.UPLOAD_DIR / str(report_id) / "originals" / doc.stored_filename
    if not original_path.exists():
        raise HTTPException(404, "DOCX file not found on disk")

    try:
        from docx import Document as DocxDocument
        docx_doc = DocxDocument(str(original_path))

        paragraphs = []
        for para in docx_doc.paragraphs:
            runs = []
            for run in para.runs:
                runs.append(DocxRun(
                    text=run.text,
                    bold=run.bold,
                    italic=run.italic,
                ))
            paragraphs.append(DocxParagraph(
                text=para.text,
                style=para.style.name if para.style else None,
                runs=runs,
            ))

        return DocxContentResponse(is_docx=True, paragraphs=paragraphs)

    except Exception as e:
        logger.error(f"Failed to read DOCX {doc.original_filename}: {e}")
        raise HTTPException(500, f"Failed to read DOCX: {str(e)}")


@app.put("/api/reports/{report_id}/documents/{doc_id}/docx-content")
async def update_docx_content(
    report_id: int,
    doc_id: int,
    request: DocxContentUpdateRequest,
    db: Session = Depends(get_db),
):
    """Update DOCX content and re-convert to PDF."""
    doc = db.query(Document).filter(Document.id == doc_id, Document.report_id == report_id).first()
    if not doc:
        raise HTTPException(404, "Document not found")

    if not doc.stored_filename.lower().endswith((".docx",)):
        raise HTTPException(400, "Not a DOCX document")

    original_path = settings.UPLOAD_DIR / str(report_id) / "originals" / doc.stored_filename
    if not original_path.exists():
        raise HTTPException(404, "DOCX file not found on disk")

    try:
        from docx import Document as DocxDocument
        docx_doc = DocxDocument(str(original_path))

        # Walk paragraphs in parallel and rewrite runs
        doc_paras = docx_doc.paragraphs
        req_paras = request.paragraphs

        for i, req_para in enumerate(req_paras):
            if i >= len(doc_paras):
                break

            orig_para = doc_paras[i]

            # If runs match count, update in place to preserve formatting
            if len(req_para.runs) == len(orig_para.runs):
                for j, req_run in enumerate(req_para.runs):
                    orig_para.runs[j].text = req_run.text
                    if req_run.bold is not None:
                        orig_para.runs[j].bold = req_run.bold
                    if req_run.italic is not None:
                        orig_para.runs[j].italic = req_run.italic
            else:
                # Clear and rewrite — preserves paragraph-level formatting
                for run in orig_para.runs:
                    run.text = ""
                if orig_para.runs:
                    orig_para.runs[0].text = req_para.text
                elif req_para.text:
                    orig_para.add_run(req_para.text)

        docx_doc.save(str(original_path))

        # Re-convert to PDF
        pdf_dir = settings.UPLOAD_DIR / str(report_id) / "pdfs"
        pdf_dir.mkdir(parents=True, exist_ok=True)

        # Clear page cache since the doc changed
        cache_dir = settings.UPLOAD_DIR / str(report_id) / "output" / "page_cache"
        if cache_dir.exists():
            shutil.rmtree(cache_dir)

        pdf_path = await async_convert_to_pdf(original_path, pdf_dir)
        if pdf_path:
            doc.pdf_filename = pdf_path.name
            page_count = await async_get_pdf_page_count(pdf_path)
            if page_count is not None:
                doc.page_count = page_count
            db.commit()
            return {"status": "ok", "page_count": doc.page_count}
        else:
            raise HTTPException(500, "PDF conversion failed after DOCX update")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to update DOCX {doc.original_filename}: {e}")
        raise HTTPException(500, f"DOCX update failed: {str(e)}")


# ---- Auto-Split Download ----

@app.get("/api/reports/{report_id}/download-auto")
def download_auto(report_id: int, db: Session = Depends(get_db)):
    """Smart download: single PDF if < 20MB, zip of split parts if larger."""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report or not report.assembled_filename:
        raise HTTPException(404, "No assembled report")

    output_dir = settings.UPLOAD_DIR / str(report_id) / "output"
    pdf_path = output_dir / report.assembled_filename
    if not pdf_path.exists():
        raise HTTPException(404, "Assembled PDF not found")

    file_size = pdf_path.stat().st_size
    threshold = 20 * 1024 * 1024  # 20 MB

    if file_size <= threshold:
        return FileResponse(pdf_path, media_type="application/pdf", filename=pdf_path.name)

    # Split and zip
    from splitter import split_pdf
    parts = split_pdf(pdf_path, 20.0)

    if len(parts) <= 1:
        return FileResponse(pdf_path, media_type="application/pdf", filename=pdf_path.name)

    # Create zip in memory
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for part in parts:
            part_path = Path(part["path"])
            if part_path.exists():
                zf.write(part_path, part_path.name)
    zip_buffer.seek(0)

    zip_filename = f"{Path(report.assembled_filename).stem}_split.zip"
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{zip_filename}"'},
    )


# ---- Helpers ----

def _report_to_response(report: Report) -> dict:
    pipeline_duration = None
    if report.pipeline_start_time and report.pipeline_end_time:
        delta = report.pipeline_end_time - report.pipeline_start_time
        pipeline_duration = int(delta.total_seconds())
    return {
        "id": report.id,
        "name": report.name,
        "address": report.address,
        "project_number": report.project_number,
        "has_reliance_letter": report.has_reliance_letter,
        "status": report.status,
        "document_count": report.document_count,
        "assembled_filename": report.assembled_filename,
        "assembled_size": report.assembled_size,
        "compressed_size": report.compressed_size,
        "pipeline_duration": pipeline_duration,
        "created_at": report.created_at,
        "updated_at": report.updated_at,
    }


def _doc_to_response(doc: Document) -> dict:
    return {
        "id": doc.id,
        "report_id": doc.report_id,
        "original_filename": doc.original_filename,
        "original_path": doc.original_path,
        "stored_filename": doc.stored_filename,
        "pdf_filename": doc.pdf_filename,
        "file_size": doc.file_size,
        "page_count": doc.page_count,
        "category": doc.category,
        "subcategory": doc.subcategory,
        "confidence": doc.confidence,
        "reasoning": doc.reasoning,
        "sort_order": doc.sort_order,
        "status": doc.status,
        "is_included": doc.is_included,
        "has_docx_source": doc.stored_filename.lower().endswith(".docx"),
        "created_at": doc.created_at,
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)
