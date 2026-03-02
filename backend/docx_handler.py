"""DOCX file reading, editing, and preview handling."""

import io
import json
import logging
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


def read_docx_content(file_path: Path) -> dict:
    """
    Read DOCX file and extract all content as editable blocks.
    
    Returns:
        {
            "blocks": [
                {"id": "p0", "type": "paragraph", "text": "...", "style": "Normal"},
                {"id": "t0", "type": "table", "rows": [...], "cols": 3},
            ],
            "metadata": {...}
        }
    """
    try:
        doc = Document(str(file_path))
        blocks = []
        block_id = 0
        
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                para = element.getparent().find(element.tag)
                text = ''.join(node.text for node in para.itertext() if node.text)
                style = para.style.name if para.style else "Normal"
                
                blocks.append({
                    "id": f"p{block_id}",
                    "type": "paragraph",
                    "text": text,
                    "style": style,
                })
                block_id += 1
                
            elif element.tag.endswith('tbl'):  # Table
                rows = []
                for row in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr'):
                    cols = []
                    for cell in row.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc'):
                        cell_text = ''.join(node.text for node in cell.itertext() if node.text)
                        cols.append(cell_text)
                    rows.append(cols)
                
                if rows:
                    blocks.append({
                        "id": f"t{block_id}",
                        "type": "table",
                        "rows": rows,
                        "cols": max(len(r) for r in rows) if rows else 0,
                    })
                    block_id += 1
        
        return {
            "status": "ok",
            "blocks": blocks,
            "total_blocks": len(blocks),
            "filename": file_path.name,
        }
    except Exception as e:
        logger.error(f"Error reading DOCX: {e}")
        return {"status": "error", "message": str(e)}


def update_docx_content(file_path: Path, changes: list[dict]) -> dict:
    """
    Apply text changes to DOCX file.
    
    Changes format:
    [
        {"id": "p0", "text": "new text"},
        {"id": "p1", "text": "another change"},
    ]
    """
    try:
        doc = Document(str(file_path))
        
        # Build map of block IDs to document elements
        block_map = {}
        block_id = 0
        
        for element in doc.element.body:
            if element.tag.endswith('p'):
                block_map[f"p{block_id}"] = element
                block_id += 1
            elif element.tag.endswith('tbl'):
                block_map[f"t{block_id}"] = element
                block_id += 1
        
        # Apply changes
        updated = 0
        for change in changes:
            block_id = change.get("id")
            new_text = change.get("text", "")
            
            if block_id not in block_map:
                continue
            
            element = block_map[block_id]
            
            # Clear existing content and set new text
            if element.tag.endswith('p'):
                # Paragraph
                para = element
                for run in para.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
                    para.remove(run)
                
                # Add new text
                para.text = new_text
                updated += 1
        
        # Save modified document
        doc.save(str(file_path))
        
        return {
            "status": "ok",
            "updated": updated,
            "filename": file_path.name,
        }
    except Exception as e:
        logger.error(f"Error updating DOCX: {e}")
        return {"status": "error", "message": str(e)}


def docx_to_html(file_path: Path) -> str:
    """
    Convert DOCX to HTML for preview in browser.
    """
    try:
        doc = Document(str(file_path))
        html_parts = ['<div class="docx-preview">']
        
        for para in doc.paragraphs:
            text = para.text
            style = para.style.name if para.style else "Normal"
            
            # Basic style mapping
            class_name = "docx-para"
            if "Heading" in style:
                level = int(style.split()[-1]) if style[-1].isdigit() else 1
                html_parts.append(f'<h{level} class="{class_name}">{text}</h{level}>')
            else:
                html_parts.append(f'<p class="{class_name}">{text}</p>')
        
        # Tables
        for table in doc.tables:
            html_parts.append('<table class="docx-table">')
            for row in table.rows:
                html_parts.append('<tr>')
                for cell in row.cells:
                    html_parts.append(f'<td>{cell.text}</td>')
                html_parts.append('</tr>')
            html_parts.append('</table>')
        
        html_parts.append('</div>')
        return '\n'.join(html_parts)
    except Exception as e:
        logger.error(f"Error converting DOCX to HTML: {e}")
        return f"<p>Error reading document: {e}</p>"


def create_docx_from_text(text: str, title: str = "Document") -> bytes:
    """
    Create a DOCX file from plain text.
    
    Returns bytes that can be saved or sent to client.
    """
    try:
        doc = Document()
        doc.add_heading(title, level=1)
        
        # Add paragraphs, respecting line breaks
        for paragraph in text.split('\n\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Save to bytes
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        return docx_bytes.getvalue()
    except Exception as e:
        logger.error(f"Error creating DOCX: {e}")
        raise
